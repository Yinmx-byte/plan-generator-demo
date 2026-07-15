from __future__ import annotations

import argparse
import asyncio
import json
import re
import statistics
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document

BACKEND_ROOT = Path(__file__).resolve().parents[2]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from agents.evaluation_agent import evaluate_with_agent


CANONICAL_HEADINGS = [
    "背景",
    "检修类型",
    "现场环境",
    "实施计划",
    "检修窗口",
    "实施人员",
    "风险评估",
    "影响范围",
    "危险点分析",
    "安全措施",
    "授权",
    "备份",
    "验证",
    "双人复核",
    "实施步骤",
    "检修前验证",
    "检修操作",
    "检修后验证",
    "回滚步骤",
    "回滚操作",
    "回滚后验证",
]

RISK_ITEMS = ["影响范围", "危险点分析", "安全措施", "授权", "备份", "验证", "双人复核"]
OPERATION_ITEMS = ["备份", "检修前验证", "检修操作", "检修后验证"]
ROLLBACK_ITEMS = ["回滚操作", "回滚后验证"]
GENERIC_PHRASES = ["根据实际情况", "视情况", "按需", "相关人员", "确保正常", "进行检查", "完成后验证"]
ACTION_VERBS = [
    "登录", "进入", "选择", "定位", "单击", "点击", "填写", "输入", "配置", "创建", "修改",
    "删除", "回收", "释放", "重启", "核对", "确认", "验证", "记录", "截图", "导出", "执行",
]

COMPONENT_WEIGHTS = {
    "deterministic": 0.40,
    "reference_comparison": 0.30,
    "model_review": 0.30,
}


def get_quality_rule_catalog() -> dict[str, Any]:
    return {
        "writeback": {
            "title": "Skill 写回规则",
            "description": "只根据实际生成 DOCX 的缺陷形成候选修改，不直接覆盖源 Skill。",
            "items": [
                {"title": "以生成结果为依据", "detail": "仅使用候选 DOCX 的确定性检查、同类优质文档对比和独立模型评审结果。"},
                {"title": "评估器与产品 Skill 解耦", "detail": "被评估 Skill 只负责生成方案，不提供自己的评分标准，避免循环自证。"},
                {"title": "缺陷反推通用规则", "detail": "只把可复用的结构、风险、步骤与回滚约束写入候选版本，不固化人员、实例、时间等样例值。"},
                {"title": "人工确认与版本回退", "detail": "候选修改先展示差异，用户确认后才应用，并保留原版本快照。"},
            ],
        },
        "scoring": {
            "title": "文档评分规则",
            "description": "总分由三类独立证据加权：确定性规则 40%、高质量文档对比 30%、模型评审 30%。",
            "items": [
                {"title": "确定性规则（40%）", "detail": "Python 稳定检查固定章节、人员、实施步骤、回滚闭环、表格和编号格式。"},
                {"title": "高质量文档对比（30%）", "detail": "按产品、动作和网络环境从远程参考库匹配同类 DOCX，对比结构、步骤粒度、风险覆盖和格式。"},
                {"title": "独立模型评审（30%）", "detail": "独立 ReActAgent 判断可执行性、空话、需求失真、风险回滚相关性和参考资料个性化信息污染。"},
                {"title": "评分不依赖源 Skill", "detail": "源 Skill 仅用于定位待优化文件，不作为评分依据。"},
            ],
            "generic_phrases": GENERIC_PHRASES,
        },
    }


@dataclass
class DocProfile:
    path: str
    headings: list[str]
    paragraphs: list[str]
    table_count: int

    @property
    def text(self) -> str:
        return "\n".join(self.paragraphs)

    @property
    def operation_lines(self) -> list[str]:
        return [
            text for text in self.paragraphs
            if any(verb in text for verb in ACTION_VERBS) and len(normalize(text)) >= 12
        ]


def normalize(text: str) -> str:
    return re.sub(r"\s+", "", text or "").lower()


def normalize_heading(text: str) -> str:
    value = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", text.strip())
    value = re.sub(r"^\d+(?:\.\d+)*[、.．]?\s*", "", value)
    return normalize(value)


def read_docx(path: Path) -> DocProfile:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                paragraphs.append(text)
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if (
            style.startswith("Heading")
            or re.match(r"^[一二三四五六七八九十]+、", text)
            or re.match(r"^\d+(?:\.\d+)+\s*", text)
            or normalize_heading(text) in {normalize(item) for item in CANONICAL_HEADINGS}
        ):
            headings.append(text)
    return DocProfile(str(path), headings, paragraphs, len(doc.tables))


def read_references(reference_dir: Path) -> list[DocProfile]:
    references = []
    for path in sorted(reference_dir.glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        try:
            references.append(read_docx(path))
        except Exception:
            continue
    if not references:
        raise RuntimeError(f"优质方案参考目录中没有可读取的 DOCX：{reference_dir}")
    return references


def coverage(items: Iterable[str], text: str) -> tuple[int, list[str]]:
    values = list(items)
    normalized_text = normalize(text)
    missing = [item for item in values if normalize(item) not in normalized_text]
    score = round((len(values) - len(missing)) * 100 / len(values)) if values else 100
    return score, missing


def weighted_score(scores: dict[str, int], weights: dict[str, float]) -> int:
    return round(sum(scores[key] * weights[key] for key in weights))


def deterministic_evaluation(candidate: DocProfile, state: dict[str, Any]) -> dict[str, Any]:
    structure_score, structure_missing = coverage(CANONICAL_HEADINGS, candidate.text)
    risk_score, risk_missing = coverage(RISK_ITEMS, candidate.text)
    operation_score, operation_missing = coverage(OPERATION_ITEMS, candidate.text)
    rollback_score, rollback_missing = coverage(ROLLBACK_ITEMS, candidate.text)

    personnel_values = [
        str(state.get(key, "")).strip()
        for key in ("provider", "executor", "reviewer", "security_officer")
        if str(state.get(key, "")).strip()
    ]
    personnel_score, personnel_missing = coverage(personnel_values, candidate.text) if personnel_values else (100, [])
    duplicate_numbering = bool(re.search(r"\b\d+、\s*\d+、", candidate.text))
    format_score = min(100, 50 + candidate.table_count * 20)
    if duplicate_numbering:
        format_score = min(format_score, 70)
    generic_hits = [phrase for phrase in GENERIC_PHRASES if phrase in candidate.text]
    if generic_hits:
        operation_score = max(0, operation_score - min(30, len(generic_hits) * 6))

    scores = {
        "structure": structure_score,
        "personnel": personnel_score,
        "risk": risk_score,
        "operation": operation_score,
        "rollback": rollback_score,
        "format": format_score,
    }
    score = weighted_score(
        scores,
        {"structure": 0.25, "personnel": 0.15, "risk": 0.20, "operation": 0.20, "rollback": 0.15, "format": 0.05},
    )
    findings: list[dict[str, str]] = []
    missing_groups = {
        "固定章节": structure_missing,
        "风险检查项": risk_missing,
        "实施链路": operation_missing,
        "回滚闭环": rollback_missing,
        "人员": personnel_missing,
    }
    for label, missing in missing_groups.items():
        if missing:
            findings.append({
                "severity": "high" if label in {"实施链路", "回滚闭环"} else "medium",
                "category": "deterministic",
                "message": f"确定性检查发现{label}缺失：" + "、".join(missing[:10]),
                "suggested_skill_change": f"在生成规则中明确要求完整输出{label}，并在交付前逐项自检。",
            })
    if generic_hits:
        findings.append({
            "severity": "medium",
            "category": "deterministic",
            "message": "实施内容包含空泛措辞：" + "、".join(generic_hits),
            "suggested_skill_change": "禁止空泛措辞，要求每步写明入口、对象、参数、动作、预期结果和留痕。",
        })
    if duplicate_numbering:
        findings.append({
            "severity": "medium",
            "category": "deterministic",
            "message": "候选文档存在重复编号。",
            "suggested_skill_change": "要求列表项正文不自带序号，由渲染工具统一生成编号。",
        })
    return {"score": score, "dimension_scores": scores, "findings": findings}


def reference_evaluation(candidate: DocProfile, references: list[DocProfile]) -> dict[str, Any]:
    reference_heading_sets = [{normalize_heading(item) for item in doc.headings} for doc in references]
    consensus: set[str] = set()
    threshold = max(1, (len(references) + 1) // 2)
    all_headings = set().union(*reference_heading_sets)
    for heading in all_headings:
        if sum(heading in values for values in reference_heading_sets) >= threshold:
            consensus.add(heading)
    candidate_headings = {normalize_heading(item) for item in candidate.headings}
    structure_score = round(100 * len(consensus & candidate_headings) / len(consensus)) if consensus else 100

    operation_counts = [max(1, len(doc.operation_lines)) for doc in references]
    operation_lengths = [
        statistics.mean(len(normalize(line)) for line in doc.operation_lines)
        for doc in references if doc.operation_lines
    ]
    target_count = statistics.median(operation_counts)
    target_length = statistics.median(operation_lengths) if operation_lengths else 20
    candidate_count = len(candidate.operation_lines)
    candidate_length = statistics.mean(len(normalize(line)) for line in candidate.operation_lines) if candidate.operation_lines else 0
    operation_score = round(
        100 * (0.6 * min(1, candidate_count / target_count) + 0.4 * min(1, candidate_length / target_length))
    )

    risk_scores = [coverage(RISK_ITEMS, doc.text)[0] for doc in references]
    candidate_risk = coverage(RISK_ITEMS, candidate.text)[0]
    target_risk = max(1, statistics.median(risk_scores))
    risk_score = round(100 * min(1, candidate_risk / target_risk))

    target_tables = max(1, statistics.median([max(1, doc.table_count) for doc in references]))
    table_ratio = min(candidate.table_count, target_tables) / max(candidate.table_count, target_tables)
    target_paragraphs = max(1, statistics.median([len(doc.paragraphs) for doc in references]))
    paragraph_ratio = min(len(candidate.paragraphs), target_paragraphs) / max(len(candidate.paragraphs), target_paragraphs)
    format_score = round(100 * (0.6 * table_ratio + 0.4 * paragraph_ratio))

    scores = {
        "reference_structure": structure_score,
        "reference_operation_granularity": operation_score,
        "reference_risk_coverage": risk_score,
        "reference_format": format_score,
    }
    score = weighted_score(
        scores,
        {"reference_structure": 0.35, "reference_operation_granularity": 0.35, "reference_risk_coverage": 0.20, "reference_format": 0.10},
    )
    findings = []
    if structure_score < 80:
        findings.append({
            "severity": "medium",
            "category": "reference_comparison",
            "message": f"候选文档对同类优质方案稳定章节的覆盖率仅为 {structure_score}%。",
            "suggested_skill_change": "补充同类方案中稳定出现的章节，但不要复制参考文档中的业务数据。",
        })
    if operation_score < 75:
        findings.append({
            "severity": "high",
            "category": "reference_comparison",
            "message": f"实施步骤粒度低于同类优质方案，当前对比得分 {operation_score}。",
            "suggested_skill_change": "要求操作步骤达到同类方案的动作数量和细节密度，写明控制台入口、目标对象、参数、动作和预期结果。",
        })
    if risk_score < 75:
        findings.append({
            "severity": "medium",
            "category": "reference_comparison",
            "message": f"风险覆盖低于同类优质方案，当前对比得分 {risk_score}。",
            "suggested_skill_change": "按本次动作补齐影响范围、危险点、安全措施和验证留痕。",
        })
    return {"score": score, "dimension_scores": scores, "findings": findings}


def build_reference_summaries(references: list[DocProfile]) -> list[dict[str, Any]]:
    return [
        {
            "filename": Path(doc.path).name,
            "headings": doc.headings[:30],
            "table_count": doc.table_count,
            "operation_examples": doc.operation_lines[:12],
        }
        for doc in references
    ]


def load_state(raw: str) -> dict[str, Any]:
    if not raw:
        return {}
    value = json.loads(raw)
    if not isinstance(value, dict):
        raise ValueError("state JSON 必须是对象")
    return value


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference-dir", required=True)
    parser.add_argument("--candidate-docx", required=True)
    parser.add_argument("--state-json", default="{}")
    parser.add_argument("--output")
    args = parser.parse_args()

    references = read_references(Path(args.reference_dir))
    candidate = read_docx(Path(args.candidate_docx))
    state = load_state(args.state_json)
    deterministic = deterministic_evaluation(candidate, state)
    reference = reference_evaluation(candidate, references)
    model = asyncio.run(
        evaluate_with_agent(
            candidate_text=candidate.text,
            state=state,
            references=build_reference_summaries(references),
        )
    )
    component_scores = {
        "deterministic": deterministic["score"],
        "reference_comparison": reference["score"],
        "model_review": model["score"],
    }
    score = weighted_score(component_scores, COMPONENT_WEIGHTS)
    findings = deterministic["findings"] + reference["findings"] + model["findings"]
    result = {
        "evaluation_mode": "generated_docx",
        "score": score,
        "score_weights": COMPONENT_WEIGHTS,
        "component_scores": component_scores,
        "dimension_scores": {
            **deterministic["dimension_scores"],
            **reference["dimension_scores"],
            **model["dimension_scores"],
        },
        "reference_summary": {
            "count": len(references),
            "files": [Path(doc.path).name for doc in references],
        },
        "candidate": {
            "type": "docx",
            "path": candidate.path,
            "headings": candidate.headings,
            "tables": candidate.table_count,
        },
        "findings": findings,
        "model_review_summary": model["summary"],
        "recommended_patch_summary": (
            "生成文档存在质量缺陷；应将问题中可复用的结构、风险、实施与回滚约束整理为 Skill 候选规则。"
            if findings
            else "三类评估均未发现明显问题，暂不建议修改源 Skill。"
        ),
    }
    output = json.dumps(result, ensure_ascii=False, indent=2)
    if args.output:
        Path(args.output).write_text(output, encoding="utf-8")
    else:
        print(output)


if __name__ == "__main__":
    main()
