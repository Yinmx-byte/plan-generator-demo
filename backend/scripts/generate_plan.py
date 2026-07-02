#!/usr/bin/env python3
"""Render maintenance-plan JSON into a Word document.

The document structure is supplied by the Skill-generated JSON. This script
only provides reusable rendering primitives.
"""

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt, RGBColor


BACKEND_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_STYLE_PATH = (
    BACKEND_ROOT
    / "skills"
    / "maintenance-plan-composer"
    / "references"
    / "document-style.json"
)
REQUIRED_STYLES = {
    "body",
    "title",
    "cover_metadata",
    "heading_1",
    "heading_2",
    "heading_3",
    "table",
    "environment_table",
    "footer_page_number",
}


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def load_style_contract(style_contract: Any = None) -> dict[str, Any]:
    """Load and validate a machine-readable document style contract."""
    if style_contract is None:
        path = DEFAULT_STYLE_PATH
        contract = json.loads(path.read_text(encoding="utf-8"))
    elif isinstance(style_contract, (str, Path)):
        path = Path(style_contract)
        contract = json.loads(path.read_text(encoding="utf-8"))
    elif isinstance(style_contract, dict):
        contract = deepcopy(style_contract)
    else:
        raise TypeError("style_contract must be a dict, path, or None")

    if not isinstance(contract.get("page"), dict):
        raise ValueError("style_contract.page must be an object")
    if not isinstance(contract.get("cover"), dict):
        raise ValueError("style_contract.cover must be an object")
    if not isinstance(contract.get("numbering"), dict):
        raise ValueError("style_contract.numbering must be an object")
    if not isinstance(contract.get("checkbox"), dict):
        raise ValueError("style_contract.checkbox must be an object")
    styles = contract.get("styles")
    if not isinstance(styles, dict):
        raise ValueError("style_contract.styles must be an object")
    missing = sorted(REQUIRED_STYLES - set(styles))
    if missing:
        raise ValueError(f"style_contract is missing styles: {', '.join(missing)}")
    for key in ("width_cm", "height_cm", "margins_cm"):
        if key not in contract["page"]:
            raise ValueError(f"style_contract.page is missing: {key}")
    for key in ("top", "bottom", "left", "right"):
        if key not in contract["page"]["margins_cm"]:
            raise ValueError(f"style_contract.page.margins_cm is missing: {key}")
    for key in (
        "logo_path",
        "logo_width_cm",
        "logo_alignment",
        "top_spacers",
        "middle_spacers",
        "title_style",
        "metadata_style",
    ):
        if key not in contract["cover"]:
            raise ValueError(f"style_contract.cover is missing: {key}")
    for style_name in REQUIRED_STYLES:
        resolved = resolve_style(contract, style_name)
        for key in ("font_name", "font_size_pt", "color", "line_spacing_pt"):
            if key not in resolved:
                raise ValueError(f"style_contract.styles.{style_name} is missing: {key}")
    border = resolve_style(contract, "table").get("border")
    if not isinstance(border, dict) or not {"style", "size", "color"} <= set(border):
        raise ValueError("style_contract.styles.table.border is incomplete")
    return contract


def resolve_style(style_contract: dict[str, Any], style_name: str) -> dict[str, Any]:
    """Resolve one named style, including optional single/multi-level inheritance."""
    styles = style_contract["styles"]
    if style_name not in styles:
        raise ValueError(f"Unknown document style: {style_name}")
    resolved: dict[str, Any] = {}
    visiting: set[str] = set()

    def merge(name: str) -> None:
        if name in visiting:
            raise ValueError(f"Circular document style inheritance: {name}")
        definition = styles.get(name)
        if not isinstance(definition, dict):
            raise ValueError(f"Document style must be an object: {name}")
        visiting.add(name)
        parent = definition.get("extends")
        if parent:
            merge(str(parent))
        resolved.update({key: value for key, value in definition.items() if key != "extends"})
        visiting.remove(name)

    merge(style_name)
    return resolved


def chars_to_cm(chars: Any, font_size_pt: Any) -> float:
    """Convert Chinese-character indentation to centimeters for Word."""
    return float(chars or 0) * float(font_size_pt) * 2.54 / 72


def rgb_color(value: Any) -> RGBColor:
    text = str(value or "000000").lstrip("#")
    if not re.fullmatch(r"[0-9a-fA-F]{6}", text):
        raise ValueError(f"Invalid RGB color: {value}")
    return RGBColor.from_string(text.upper())


def set_cell_border(cell, **kwargs):
    """Set borders for a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def add_table_borders(table, border: dict[str, Any]):
    for row in table.rows:
        for cell in row.cells:
            for edge in ["top", "left", "bottom", "right"]:
                set_cell_border(cell, **{edge: {
                    "val": border["style"],
                    "sz": border["size"],
                    "color": border["color"],
                }})


def set_run_font(run, style: dict[str, Any], *, bold: Any = None):
    font_name = str(style["font_name"])
    run.font.name = font_name
    run.font.size = Pt(float(style["font_size_pt"]))
    run.font.color.rgb = rgb_color(style.get("color"))
    run.bold = bool(style.get("bold", False) if bold is None else bold)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def apply_paragraph_style(paragraph, style: dict[str, Any], *, empty: bool = False, alignment=None):
    paragraph.paragraph_format.space_before = Pt(
        0 if empty else float(style.get("space_before_pt", 0))
    )
    paragraph.paragraph_format.space_after = Pt(float(style.get("space_after_pt", 0)))
    paragraph.paragraph_format.line_spacing = Pt(float(style["line_spacing_pt"]))
    font_size = style["font_size_pt"]
    if "first_line_indent_chars" in style:
        paragraph.paragraph_format.first_line_indent = Cm(
            chars_to_cm(style["first_line_indent_chars"], font_size)
        )
    if "left_indent_chars" in style:
        paragraph.paragraph_format.left_indent = Cm(
            chars_to_cm(style["left_indent_chars"], font_size)
        )
    effective_alignment = (
        alignment if alignment is not None else ALIGNMENTS.get(style.get("alignment"))
    )
    if effective_alignment is not None:
        paragraph.alignment = effective_alignment


def add_paragraph(
    doc,
    style_contract,
    text="",
    style_name="body",
    bold=False,
    alignment=None,
    style_overrides=None,
):
    style = resolve_style(style_contract, style_name)
    if style_overrides:
        style.update(style_overrides)
    p = doc.add_paragraph()
    apply_paragraph_style(p, style, empty=not text, alignment=alignment)
    run = p.add_run(str(text))
    set_run_font(run, style, bold=bold or None)
    return p


def add_heading_text(doc, style_contract, text, level=1):
    style_name = "title" if level == 0 else f"heading_{level}"
    style = resolve_style(style_contract, style_name)
    style_name = f"Heading {level}" if level in {1, 2, 3} else None
    p = doc.add_paragraph(style=style_name)
    apply_paragraph_style(p, style)
    run = p.add_run(str(text))
    set_run_font(run, style)
    return p


def apply_document_defaults(doc, style_contract):
    body_style = resolve_style(style_contract, "body")
    style = doc.styles["Normal"]
    font = style.font
    font.name = body_style["font_name"]
    font.size = Pt(float(body_style["font_size_pt"]))
    font.color.rgb = rgb_color(body_style.get("color"))
    style.element.rPr.rFonts.set(qn("w:eastAsia"), body_style["font_name"])
    apply_paragraph_style(style, body_style)

    for level in (1, 2, 3):
        heading = resolve_style(style_contract, f"heading_{level}")
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = heading["font_name"]
        heading_style.font.size = Pt(float(heading["font_size_pt"]))
        heading_style.font.bold = bool(heading.get("bold", False))
        heading_style.font.color.rgb = rgb_color(heading.get("color"))
        heading_style.element.rPr.rFonts.set(qn("w:eastAsia"), heading["font_name"])
        apply_paragraph_style(heading_style, heading)

    page = style_contract["page"]
    margins = page["margins_cm"]
    for section in doc.sections:
        section.page_width = Cm(float(page["width_cm"]))
        section.page_height = Cm(float(page["height_cm"]))
        section.top_margin = Cm(float(margins["top"]))
        section.bottom_margin = Cm(float(margins["bottom"]))
        section.left_margin = Cm(float(margins["left"]))
        section.right_margin = Cm(float(margins["right"]))
        add_page_number(section, style_contract)


def add_page_number(section, style_contract):
    """Add a centered Chinese-style page number such as —1—."""
    style = resolve_style(style_contract, "footer_page_number")
    paragraph = section.footer.paragraphs[0]
    apply_paragraph_style(paragraph, style)

    prefix = paragraph.add_run(str(style.get("prefix", "")))
    set_run_font(prefix, style)
    page_run = paragraph.add_run()
    set_run_font(page_run, style)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached_value = OxmlElement("w:t")
    cached_value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([begin, instruction, separate, cached_value, end])
    suffix = paragraph.add_run(str(style.get("suffix", "")))
    set_run_font(suffix, style)


def add_simple_table(doc, style_contract: dict[str, Any], block: dict[str, Any]):
    columns = block.get("columns", [])
    rows = block.get("rows", [])
    if not columns:
        return

    table = doc.add_table(rows=len(rows) + 1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    style = resolve_style(style_contract, str(block.get("style", "table")))
    add_table_borders(table, style["border"])

    for col_idx, column in enumerate(columns):
        cell = table.rows[0].cells[col_idx]
        cell.text = str(column.get("label", column.get("key", "")))
        for paragraph in cell.paragraphs:
            apply_paragraph_style(paragraph, style)
            for run in paragraph.runs:
                set_run_font(run, style, bold=True)

    for row_idx, row in enumerate(rows, 1):
        for col_idx, column in enumerate(columns):
            key = column.get("key", column.get("label", ""))
            value = row.get(key, "") if isinstance(row, dict) else ""
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                apply_paragraph_style(paragraph, style)
                for run in paragraph.runs:
                    set_run_font(run, style)


def add_checkbox_group(doc, style_contract: dict[str, Any], block: dict[str, Any]):
    items = block.get("items", [])
    per_line = int(block.get("per_line", 2))
    row_lengths = block.get("row_lengths") or []
    checkbox = style_contract["checkbox"]
    rendered = []
    for item in items:
        label = item.get("label", "")
        checked = item.get("checked", False)
        extra = item.get("extra", "")
        mark = checkbox["checked_symbol"] if checked else checkbox["unchecked_symbol"]
        rendered.append(
            f"{mark}{checkbox['label_gap']}{label}"
            f"{(checkbox['extra_gap'] + extra) if extra else ''}"
        )

    if row_lengths:
        offset = 0
        for row_length in row_lengths:
            row = rendered[offset : offset + int(row_length)]
            if row:
                add_paragraph(doc, style_contract, checkbox["item_separator"].join(row))
            offset += int(row_length)
        return
    for idx in range(0, len(rendered), per_line):
        add_paragraph(
            doc,
            style_contract,
            checkbox["item_separator"].join(rendered[idx : idx + per_line]),
        )


def strip_manual_list_number(text: Any) -> str:
    """Remove model-supplied numbering before renderer adds its own."""
    value = str(text or "").strip()
    return re.sub(r"^\s*(?:\d+[\u3001、.,，．]|[（(]\d+[）)]|[一二三四五六七八九十]+[\u3001、])\s*", "", value)


def render_block(doc, style_contract: dict[str, Any], block: dict[str, Any]):
    block_type = block.get("type", "paragraph")

    if block_type == "heading":
        add_heading_text(
            doc,
            style_contract,
            block.get("text", ""),
            level=int(block.get("level", 1)),
        )
    elif block_type == "paragraph":
        add_paragraph(
            doc,
            style_contract,
            block.get("text", ""),
            style_name=str(block.get("style", "body")),
            bold=bool(block.get("bold", False)),
            alignment=ALIGNMENTS.get(block.get("align")),
        )
    elif block_type == "paragraphs":
        for text in block.get("items", []):
            add_paragraph(doc, style_contract, text)
    elif block_type == "numbered_list":
        for idx, text in enumerate(block.get("items", []), 1):
            add_paragraph(doc, style_contract, f"{idx}、{strip_manual_list_number(text)}")
    elif block_type == "plain_list":
        prefix = block.get("prefix", "")
        for text in block.get("items", []):
            add_paragraph(doc, style_contract, f"{prefix}{text}")
    elif block_type == "key_values":
        for item in block.get("items", []):
            add_paragraph(
                doc,
                style_contract,
                f'{item.get("label", "")}{item.get("separator", "：")}{item.get("value", "")}',
            )
    elif block_type == "checkbox_group":
        add_checkbox_group(doc, style_contract, block)
    elif block_type == "table":
        add_simple_table(doc, style_contract, block)
    elif block_type == "spacer":
        add_paragraph(doc, style_contract, "")


def normalize_text_items(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [line.strip() for line in value.splitlines() if line.strip()]
    if isinstance(value, list):
        items = []
        for item in value:
            if isinstance(item, str) and item.strip():
                items.append(item.strip())
            elif isinstance(item, dict):
                text = item.get("text") or item.get("content") or item.get("description")
                if text:
                    items.extend(normalize_text_items(text))
        return items
    return [str(value).strip()] if str(value).strip() else []


def normalize_block(block: Any) -> list[dict[str, Any]]:
    if isinstance(block, str):
        return [{"type": "paragraph", "text": block}]
    if not isinstance(block, dict):
        return []
    if "type" in block:
        return [block]
    if "columns" in block and "rows" in block:
        return [{"type": "table", **block}]
    if "items" in block:
        return [{"type": "paragraphs", "items": normalize_text_items(block.get("items"))}]
    text = block.get("text") or block.get("content") or block.get("description")
    if text:
        return [{"type": "paragraph", "text": text}]
    return []


def normalize_section_blocks(section: dict[str, Any]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for block in section.get("blocks") or []:
        blocks.extend(normalize_block(block))

    if not blocks:
        for key in ("content", "body", "text", "description"):
            for text in normalize_text_items(section.get(key)):
                blocks.append({"type": "paragraph", "text": text})
        for key in ("items", "points", "steps"):
            items = normalize_text_items(section.get(key))
            if items:
                blocks.append({"type": "numbered_list" if key == "steps" else "paragraphs", "items": items})

    children = section.get("children") or section.get("subsections") or []
    for child in children:
        if isinstance(child, str):
            blocks.append({"type": "paragraph", "text": child})
            continue
        if not isinstance(child, dict):
            continue
        child_title = child.get("heading") or child.get("title")
        if child_title:
            blocks.append({"type": "heading", "text": child_title, "level": int(child.get("level", 2))})
        blocks.extend(normalize_section_blocks(child))
    return blocks


def numbered_section_heading(style_contract: dict[str, Any], text: Any, index: int) -> str:
    heading = str(text or "").strip()
    if not heading:
        return heading
    numbering = style_contract["numbering"]
    suffix = str(numbering["level_1_suffix"])
    if suffix in heading[:3]:
        return heading
    numerals = numbering["level_1_numerals"]
    if index < len(numerals):
        return f"{numerals[index]}{suffix}{heading}"
    return heading


def resolve_logo_path(style_contract: dict[str, Any]) -> Path | None:
    configured = style_contract["cover"].get("logo_path")
    if not configured:
        return None
    path = Path(str(configured))
    if path.is_absolute():
        raise ValueError("cover.logo_path must be relative to backend/assets")
    resolved = (BACKEND_ROOT / path).resolve()
    assets_root = (BACKEND_ROOT / "assets").resolve()
    if not resolved.is_relative_to(assets_root):
        raise ValueError("cover.logo_path must stay within backend/assets")
    return resolved


def render_document(doc, spec: dict[str, Any], style_contract: dict[str, Any]):
    title = spec.get("title", "检修方案")
    cover = spec.get("cover", {})
    if cover is not False:
        cover_style = style_contract["cover"]
        logo_path = resolve_logo_path(style_contract)
        if logo_path and logo_path.exists():
            p = doc.add_paragraph()
            p.alignment = ALIGNMENTS[str(cover_style["logo_alignment"])]
            run = p.add_run()
            run.add_picture(str(logo_path), width=Cm(float(cover_style["logo_width_cm"])))
        blank_overrides = {
            "space_before_pt": 0,
            "space_after_pt": 0,
            "first_line_indent_chars": 0,
        }
        for _ in range(int(cover_style["top_spacers"])):
            add_paragraph(doc, style_contract, "", style_overrides=blank_overrides)
        add_paragraph(
            doc,
            style_contract,
            title,
            style_name=str(cover_style["title_style"]),
        )
        for _ in range(int(cover_style["middle_spacers"])):
            add_paragraph(doc, style_contract, "", style_overrides=blank_overrides)
        header = spec.get("header", [])
        for line in header:
            if isinstance(line, str):
                line = {"text": line}
            add_paragraph(
                doc,
                style_contract,
                line.get("text", ""),
                style_name=str(line.get("style", cover_style["metadata_style"])),
                alignment=ALIGNMENTS.get(line.get("align")),
            )
        doc.add_page_break()
    else:
        add_paragraph(
            doc,
            style_contract,
            title,
            style_name="title",
        )

        for line in spec.get("header", []):
            if isinstance(line, str):
                line = {"text": line}
            add_paragraph(
                doc,
                style_contract,
                line.get("text", ""),
                style_name=str(line.get("style", "cover_metadata")),
                alignment=ALIGNMENTS.get(line.get("align")),
            )

    for section_index, section in enumerate(spec.get("sections") or spec.get("chapters") or []):
        if isinstance(section, str):
            section = {"blocks": [{"type": "paragraph", "text": section}]}
        if not isinstance(section, dict):
            continue
        heading = section.get("heading") or section.get("title")
        if heading:
            level = int(section.get("level", 1))
            if level == 1:
                heading = numbered_section_heading(style_contract, heading, section_index)
            add_heading_text(doc, style_contract, heading, level=level)
        for block in normalize_section_blocks(section):
            render_block(doc, style_contract, block)

def build_document(data, style_contract=None):
    """Build a Word document from Skill-generated document instructions."""
    resolved_contract = load_style_contract(style_contract)
    doc = Document()
    apply_document_defaults(doc, resolved_contract)

    spec = data.get("document")
    if not isinstance(spec, dict):
        raise ValueError("document must be an object following the Skill rendering contract")
    render_document(doc, spec, resolved_contract)
    return doc


def main():
    parser = argparse.ArgumentParser(description="生成国网云平台检修方案 Word 文档")
    parser.add_argument("--data", required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 文件路径")
    parser.add_argument("--style", help="可选的文档格式契约 JSON；默认读取通用 Skill 配置")
    args = parser.parse_args()

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = build_document(data, args.style)
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    print(f"检修方案已生成: {args.output}")


if __name__ == "__main__":
    main()
