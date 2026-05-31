"""
国网云平台检修方案生成服务 - 基于 AgentScope 框架。

流程：用户输入 → Skill 路由/展开 → AnthropicChatModel → 结构化 JSON → generate_plan.py → Word 文档

Skill 装卸：修改 backend/skills/ 目录下的 Skill，重启即生效。
"""

import json
import os
import re
import tempfile
import uuid
import zipfile
import inspect
import asyncio
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from agentscope.agent import ReActAgent
from agentscope.formatter import AnthropicChatFormatter, OpenAIChatFormatter
from agentscope.memory import InMemoryMemory
from agentscope.message import Msg, TextBlock
from agentscope.mcp import HttpStatefulClient, HttpStatelessClient, StdIOStatefulClient
from agentscope.model import AnthropicChatModel, OpenAIChatModel
from agentscope.tool import Toolkit, ToolResponse
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from rag import get_knowledge_base, reset_knowledge_base
from skills_runtime import SkillRegistry
ROOT = Path(__file__).parent
SKILLS_ROOT = ROOT / "skills"

from scripts.generate_plan import build_document

load_dotenv(ROOT / ".env", override=True)

# ── Skill 注册 ──────────────────────────────────────────────────
_skill_registry: Optional[SkillRegistry] = None
_toolkit: Optional[Toolkit] = None
_mcp_clients: list[Any] = []
_rag_enabled: bool = False
_chat_sessions: dict[str, dict[str, Any]] = {}
_generated_files: dict[str, Path] = {}

REQUIRED_FIELDS = {
    "background": "检修背景/检修事项",
    "maintenance_type": "检修类型",
    "network": "内外网环境",
    "location": "实施地点",
    "instances": "涉及的组件实例、组织、资源集",
    "schedule_start": "检修开始时间",
    "schedule_end": "检修结束时间",
    "provider": "方案提供人",
    "executor": "检修执行人",
    "reviewer": "检修复核人",
    "security_officer": "安全责任人",
    "ascm_account": "ASCM 授权账号",
    "bastion_account": "堡垒机账号",
}

FORM_FIELDS = [
    "background",
    "maintenance_type",
    "network",
    "location",
    "instances",
    "schedule_year",
    "schedule_start",
    "schedule_end",
    "provider",
    "executor",
    "reviewer",
    "security_officer",
    "ascm_account",
    "bastion_account",
    "ops_detail",
    "tech_params",
]


class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    execute_validation: bool = False


class ChatResetRequest(BaseModel):
    session_id: str


class PageAgentTaskRequest(BaseModel):
    task: str


def get_skill_registry() -> SkillRegistry:
    """Load Skill metadata without expanding all Skill bodies."""
    global _skill_registry
    if _skill_registry is None:
        _skill_registry = SkillRegistry(SKILLS_ROOT)
    return _skill_registry


async def reset_skill_runtime() -> None:
    """Reload skill metadata/toolkit after skill files change."""
    global _skill_registry, _toolkit
    await close_mcp_clients()
    _skill_registry = None
    _toolkit = None
    reset_knowledge_base()


async def read_file(file_path: str) -> ToolResponse:
    """读取指定 Skill 文件内容。

    Args:
        file_path: 要读取的 Skill 文件路径，通常是某个 SKILL.md。
    """
    requested = Path(file_path)
    if not requested.is_absolute():
        requested = ROOT / requested
    resolved = requested.resolve()
    skills_root = SKILLS_ROOT.resolve()
    if not str(resolved).startswith(str(skills_root)):
        raise ValueError("只能读取 backend/skills 目录下的 Skill 文件。")
    if not resolved.is_file():
        raise FileNotFoundError(str(resolved))
    content = resolved.read_text(encoding="utf-8")
    return ToolResponse(content=[TextBlock(type="text", text=content)])


def load_mcp_server_configs() -> list[dict[str, Any]]:
    config_path = ROOT / "mcp_servers.json"
    if not config_path.exists():
        return []
    config = json.loads(config_path.read_text(encoding="utf-8"))
    servers = config.get("servers", [])
    if not isinstance(servers, list):
        raise RuntimeError("mcp_servers.json 中的 servers 必须是数组。")
    return servers


def resolve_backend_path(value: str | None) -> str | None:
    if not value:
        return None
    path = Path(value)
    if not path.is_absolute():
        path = ROOT / path
    return str(path.resolve())


def build_mcp_env(env_config: dict[str, Any] | None) -> dict[str, str] | None:
    if not env_config:
        return None
    env: dict[str, str] = {}
    for key, value in env_config.items():
        if isinstance(value, str) and value.startswith("$"):
            env[key] = os.getenv(value[1:], "")
        else:
            env[key] = str(value)
    return env


async def register_mcp_servers(toolkit: Toolkit) -> list[dict[str, Any]]:
    """Register configured MCP servers as AgentScope toolkit functions."""
    global _mcp_clients
    registered = []
    for item in load_mcp_server_configs():
        name = item.get("name")
        server_type = item.get("type", "http_stateless")
        if not name:
            raise RuntimeError("mcp_servers.json 中每个 server 都必须配置 name。")

        if server_type == "http_stateless":
            client = HttpStatelessClient(
                name=name,
                transport=item.get("transport", "streamable_http"),
                url=item["url"],
                headers=item.get("headers"),
                timeout=float(item.get("timeout", 30)),
                sse_read_timeout=float(item.get("sse_read_timeout", 300)),
            )
        elif server_type == "http_stateful":
            client = HttpStatefulClient(
                name=name,
                transport=item.get("transport", "streamable_http"),
                url=item["url"],
                headers=item.get("headers"),
                timeout=float(item.get("timeout", 30)),
                sse_read_timeout=float(item.get("sse_read_timeout", 300)),
            )
            await client.connect()
            _mcp_clients.append(client)
        elif server_type == "stdio":
            client = StdIOStatefulClient(
                name=name,
                command=item["command"],
                args=item.get("args"),
                env=build_mcp_env(item.get("env")),
                cwd=resolve_backend_path(item.get("cwd")),
            )
            await client.connect()
            _mcp_clients.append(client)
        else:
            raise RuntimeError(f"不支持的 MCP server type: {server_type}")

        group_name = item.get("group_name", "mcp")
        if group_name != "basic" and group_name not in toolkit.groups:
            toolkit.create_tool_group(
                group_name=group_name,
                description=f"MCP tools from {name}",
                active=True,
                notes="Use these tools only when the task needs external MCP capabilities.",
            )

        await toolkit.register_mcp_client(
            client,
            group_name=group_name,
            enable_funcs=item.get("enable_funcs"),
            disable_funcs=item.get("disable_funcs"),
            preset_kwargs_mapping=item.get("preset_kwargs_mapping"),
            namesake_strategy=item.get("namesake_strategy", "rename"),
            execution_timeout=item.get("execution_timeout"),
        )
        registered.append(
            {
                "name": name,
                "type": server_type,
                "group_name": group_name,
            }
        )
    return registered


async def close_mcp_clients() -> None:
    global _mcp_clients
    for client in _mcp_clients:
        close = getattr(client, "close", None)
        if close:
            await close(ignore_errors=True)
    _mcp_clients = []


async def get_toolkit() -> Toolkit:
    """Create the AgentScope Toolkit with tools, Skills and MCP clients."""
    global _toolkit
    if _toolkit is not None:
        return _toolkit

    toolkit = Toolkit()
    toolkit.register_tool_function(read_file)
    for skill in get_skill_registry().skills:
        toolkit.register_agent_skill(str(skill.path))
    await register_mcp_servers(toolkit)
    _toolkit = toolkit
    return _toolkit


# ── LLM 客户端 (AgentScope AnthropicChatModel) ─────────────────
_model: Optional[AnthropicChatModel | OpenAIChatModel] = None


def get_model() -> AnthropicChatModel | OpenAIChatModel:
    global _model
    if _model is None:
        provider = os.getenv("MODEL_PROVIDER", "deepseek").lower()
        model_name = os.getenv("MODEL_NAME", "deepseek-v4-pro")
        client_kwargs = {}

        if provider == "anthropic":
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                raise RuntimeError("请在 .env 中设置 ANTHROPIC_API_KEY")
            base_url = os.getenv("ANTHROPIC_BASE_URL")
            if base_url:
                client_kwargs["base_url"] = base_url
            _model = AnthropicChatModel(
                model_name=model_name,
                api_key=api_key,
                stream=False,
                max_tokens=int(os.getenv("MAX_TOKENS", "4096")),
                client_kwargs=client_kwargs if client_kwargs else None,
            )
        elif provider in {"openai", "deepseek"}:
            env_key = "DEEPSEEK_API_KEY" if provider == "deepseek" else "OPENAI_API_KEY"
            api_key = os.getenv(env_key) or os.getenv("OPENAI_API_KEY")
            if not api_key:
                raise RuntimeError(f"请在 .env 中设置 {env_key}")
            base_url = (
                os.getenv("DEEPSEEK_BASE_URL")
                if provider == "deepseek"
                else os.getenv("OPENAI_BASE_URL")
            )
            if base_url:
                client_kwargs["base_url"] = base_url
            _model = OpenAIChatModel(
                model_name=model_name,
                api_key=api_key,
                stream=False,
                client_kwargs=client_kwargs if client_kwargs else None,
                generate_kwargs={
                    "max_tokens": int(os.getenv("MAX_TOKENS", "4096")),
                    **(
                        {"extra_body": {"thinking": {"type": "disabled"}}}
                        if provider == "deepseek"
                        else {}
                    ),
                },
            )
        else:
            raise RuntimeError(f"不支持的 MODEL_PROVIDER: {provider}")
    return _model


def get_formatter():
    provider = os.getenv("MODEL_PROVIDER", "deepseek").lower()
    if provider in {"openai", "deepseek"}:
        return OpenAIChatFormatter()
    return AnthropicChatFormatter()


def build_system_prompt() -> str:
    """Build the agent system prompt; Toolkit appends compact Skill prompts."""
    return f"""你是国网云平台检修方案生成 Agent。

遵循 AgentScope Skill 的渐进式披露原则：先根据 Skill 描述判断任务需要哪些 Skill，再通过 read_file 工具读取对应 SKILL.md。

如果用户消息中包含“编排上下文”，其中的候选 Skill 是后端初筛结果，必须优先读取这些 Skill；RAG 参考资料只作为模板、案例、API 约束和风险控制依据，不得覆盖 Skill 的硬性规则。

最终只输出 JSON，不要输出解释文字。JSON 顶层必须包含 document，document.sections 决定 Word 文档结构。"""


async def get_agent_knowledge():
    """Return AgentScope knowledge object when RAG is configured."""
    knowledge_base = get_knowledge_base(SKILLS_ROOT)
    if knowledge_base is None:
        return None
    return await knowledge_base.get_knowledge()


def build_state_text(state: dict[str, str]) -> str:
    return "\n".join(
        f"{field}: {state.get(field, '')}"
        for field in FORM_FIELDS
        if state.get(field, "").strip()
    )


def select_generation_skills(state: dict[str, str]):
    registry = get_skill_registry()
    return registry.select_skills(
        state.get("maintenance_type", ""),
        build_state_text(state),
    )


def build_selected_skill_context(selected_skills) -> str:
    if not selected_skills:
        return "未命中明确候选 Skill，请根据已注册 Skill 自行判断。"
    lines = [
        "系统已根据检修类型和需求内容初筛出候选 Skill。你必须优先读取这些 Skill 的 SKILL.md；如判断不充分，再结合已注册 Skill 追加读取。",
    ]
    for skill in selected_skills:
        lines.append(
            f"- name: {skill.name}\n"
            f"  description: {skill.description}\n"
            f"  skill_dir: {skill.path}\n"
            f"  skill_file: {skill.path / 'SKILL.md'}"
        )
    return "\n".join(lines)


def build_rag_query(state: dict[str, str], selected_skills) -> str:
    skill_text = "\n".join(
        f"{skill.name}: {skill.description}" for skill in selected_skills
    )
    return f"""检修方案参考资料检索
检修类型：{state.get("maintenance_type", "")}
网络环境：{state.get("network", "")}
实施地点：{state.get("location", "")}
涉及实例：{state.get("instances", "")}
技术参数：{state.get("tech_params", "")}
补充要求：{state.get("ops_detail", "")}
候选 Skill：
{skill_text}
请检索相似内部模板、阿里云通用检修方案、风险控制、前置检查、实施步骤、回退和验证要求。"""


async def retrieve_generation_context(state: dict[str, str], selected_skills) -> list[str]:
    knowledge_base = get_knowledge_base(SKILLS_ROOT)
    if knowledge_base is None:
        return []
    try:
        return await knowledge_base.retrieve(
            build_rag_query(state, selected_skills),
            top_k=int(os.getenv("PLAN_RAG_TOP_K", os.getenv("RAG_TOP_K", "5"))),
        )
    except Exception:
        return []


def build_rag_context(rag_chunks: list[str]) -> str:
    if not rag_chunks:
        return "当前未检索到 RAG 参考资料。若后续配置 embedding API 并添加 backend/knowledge 文档，这里会注入内部模板、历史方案和阿里云通用方案片段。"
    blocks = []
    max_chars = int(os.getenv("PLAN_RAG_CONTEXT_MAX_CHARS", "6000"))
    used = 0
    for idx, chunk in enumerate(rag_chunks, start=1):
        clean = chunk.strip()
        if not clean:
            continue
        room = max_chars - used
        if room <= 0:
            break
        clean = clean[:room]
        used += len(clean)
        blocks.append(f"[RAG-{idx}]\n{clean}")
    return "\n\n".join(blocks) if blocks else "当前未检索到 RAG 参考资料。"


async def build_generation_orchestration_context(state: dict[str, str]) -> dict[str, Any]:
    selected_skills = select_generation_skills(state)
    rag_chunks = await retrieve_generation_context(state, selected_skills)
    return {
        "selected_skill_names": [skill.name for skill in selected_skills],
        "rag_enabled": get_knowledge_base(SKILLS_ROOT) is not None,
        "rag_chunks_count": len(rag_chunks),
        "prompt_context": (
            "## 编排上下文：Skill 初筛结果\n"
            f"{build_selected_skill_context(selected_skills)}\n\n"
            "## 编排上下文：RAG 参考资料\n"
            f"{build_rag_context(rag_chunks)}\n\n"
            "## 使用要求\n"
            "- Skill 是主规则来源：文档结构、必填章节、风险点、实施步骤和脚本模板优先遵循 Skill。\n"
            "- RAG 是参考依据：用于补充内部模板措辞、历史方案经验、阿里云通用方案/API 约束，不得覆盖 Skill 的硬性规则。\n"
            "- 输出 JSON 中建议包含 evidence 字段，记录 selected_skills 和 rag_chunks_count，便于后续审计。\n"
        ),
    }


def extract_json(text: Any) -> dict:
    """从 LLM 返回的文本中提取 JSON。"""
    if isinstance(text, dict) and "document" in text:
        return text
    if not isinstance(text, str):
        text = get_response_text(text)
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        if lines[0].startswith("```"):
            text = "\n".join(lines[1:])
        if text.rstrip().endswith("```"):
            text = text.rstrip()[:-3].strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        if start < 0:
            raise

        depth = 0
        in_string = False
        escape = False
        for idx in range(start, len(text)):
            char = text[idx]
            if escape:
                escape = False
                continue
            if char == "\\":
                escape = True
                continue
            if char == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    return json.loads(text[start : idx + 1])
        raise


async def repair_and_extract_json(text: str) -> dict:
    """Repair slightly malformed model JSON and parse it again."""
    try:
        return extract_json(text)
    except json.JSONDecodeError:
        repair_prompt = f"""下面是一段模型输出的检修方案 JSON，但它可能存在漏逗号、尾随逗号、代码块包裹或混入说明文字等格式问题。
请在不改变语义和字段内容的前提下，把它修复为严格合法 JSON。

要求：
- 只输出 JSON 对象
- 不要输出 markdown
- 不要解释
- 保留 document、sections、tables、steps 等原有结构

原始输出：
{text}
"""
        response = await get_model()(
            [
                {"role": "system", "content": "你是 JSON 修复器，只输出严格合法 JSON。"},
                {"role": "user", "content": repair_prompt},
            ],
        )
        try:
            return extract_json(get_response_text(response))
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=502,
                detail="模型返回的方案 JSON 无法解析，请重试或补充更明确的需求。",
            ) from exc


def write_model_output_debug(text: str, prefix: str = "plan_model_output") -> Path:
    output_dir = Path(tempfile.gettempdir()) / "plan-generator"
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    path.write_text(text or "", encoding="utf-8")
    return path


def build_fallback_plan_data(
    state: dict[str, str],
    orchestration: dict[str, Any],
    raw_text: str = "",
) -> dict[str, Any]:
    """Build a conservative document spec when model JSON cannot be parsed."""
    maintenance_type = state.get("maintenance_type") or "检修方案"
    title = f"{maintenance_type}检修方案"
    checked_type = maintenance_type.strip()
    type_names = ["配置变更", "组件升级", "组件扩缩容", "数据库变更", "日常维护（原硬件设备）", "其他"]
    checkbox_items = [
        {
            "label": name,
            "checked": name == checked_type or (name != "其他" and name in checked_type),
            "extra": checked_type if name == "其他" and checked_type not in type_names else "",
        }
        for name in type_names
    ]
    if not any(item["checked"] for item in checkbox_items):
        checkbox_items[-1]["checked"] = True
        checkbox_items[-1]["extra"] = checked_type

    schedule_text = " ".join(
        value
        for value in [
            state.get("schedule_year", ""),
            state.get("schedule_start", ""),
            "至" if state.get("schedule_start") or state.get("schedule_end") else "",
            state.get("schedule_end", ""),
        ]
        if value
    )
    ops_hint = state.get("ops_detail") or "按对应检修类型 Skill 进行前置检查、变更实施、结果验证和回滚准备。"
    tech_hint = state.get("tech_params") or "无额外技术参数。"
    text_blob = "\n".join(
        [
            state.get("background", ""),
            state.get("maintenance_type", ""),
            state.get("instances", ""),
            state.get("tech_params", ""),
            state.get("ops_detail", ""),
        ]
    ).lower()
    if "ecs" in text_blob and any(word in text_blob for word in ["创建", "新建", "申请"]):
        item_name = state.get("background") or state.get("instances") or "创建ecs实例"
        organization = "待实施前确认"
        resource_set = "待实施前确认"
        return {
            "title": f"{item_name}检修方案",
            "department": "云运营中心平台运维处",
            "date": datetime.now().strftime("%Y年%m月%d日"),
            "evidence": {
                "selected_skills": orchestration.get("selected_skill_names", []),
                "rag_enabled": orchestration.get("rag_enabled", False),
                "rag_chunks_count": orchestration.get("rag_chunks_count", 0),
                "fallback_used": True,
                "raw_output_saved": bool(raw_text),
            },
            "document": {
                "title": f"{item_name}检修方案",
                "cover": {"logo_width_cm": 3.1, "top_spacers": 7, "middle_spacers": 8},
                "header": [
                    {"text": "云运营中心平台运维处", "font_size": 14, "align": "center"},
                    {"text": datetime.now().strftime("%Y年%m月%d日"), "font_size": 12, "align": "center"},
                ],
                "sections": [
                    {
                        "heading": "背景",
                        "blocks": [
                            {"type": "paragraph", "text": item_name},
                            {"type": "paragraph", "text": "该事项项目组提报问题工单，需检修进行处理，实现问题工单闭环。"},
                        ],
                    },
                    {"heading": "检修类型", "blocks": [{"type": "checkbox_group", "items": checkbox_items, "per_line": 3}]},
                    {
                        "heading": "现场环境",
                        "blocks": [
                            {"type": "paragraph", "text": f"（1）内网环境/外网环境：{state.get('network', '')}"},
                            {"type": "paragraph", "text": f"（2）实施地点：{state.get('location', '')}"},
                            {"type": "paragraph", "text": "（3）专有云版本：v3.16"},
                            {"type": "paragraph", "text": "（4）涉及的组件实例信息："},
                            {"type": "paragraph", "text": f"1、{item_name}"},
                            {"type": "paragraph", "text": f"组织：{organization}"},
                            {"type": "paragraph", "text": f"资源集：{resource_set}"},
                            {
                                "type": "table",
                                "columns": [
                                    {"key": "cloud_env", "label": "云环境"},
                                    {"key": "instance_name", "label": "实例名称"},
                                    {"key": "disk", "label": "磁盘"},
                                    {"key": "image", "label": "自定义镜像"},
                                    {"key": "password", "label": "密码"},
                                    {"key": "vpc", "label": "VPC ID或名称"},
                                    {"key": "vswitch", "label": "Vswitch ID或名称"},
                                    {"key": "security_group", "label": "安全组"},
                                    {"key": "spec", "label": "实例规格"},
                                    {"key": "count", "label": "数量"},
                                ],
                                "rows": [
                                    {
                                        "cloud_env": state.get("network", ""),
                                        "instance_name": state.get("instances", "") or "待实施前确认",
                                        "disk": "待实施前确认",
                                        "image": "待实施前确认",
                                        "password": "按 ASCM 平台规范生成，方案不明文展示",
                                        "vpc": "待实施前确认",
                                        "vswitch": "待实施前确认",
                                        "security_group": "待实施前确认",
                                        "spec": "待实施前确认",
                                        "count": "待实施前确认",
                                    }
                                ],
                            },
                        ],
                    },
                    {
                        "heading": "实施计划",
                        "blocks": [
                            {"type": "heading", "text": "4.1 检修窗口", "level": 2},
                            {"type": "table", "columns": [{"key": "year", "label": "年份"}, {"key": "start_time", "label": "开始时间"}, {"key": "end_time", "label": "结束时间"}], "rows": [{"year": state.get("schedule_year", ""), "start_time": state.get("schedule_start", ""), "end_time": state.get("schedule_end", "")}]},
                            {"type": "heading", "text": "4.2 实施人员", "level": 2},
                            {"type": "table", "columns": [{"key": "provider", "label": "方案提供人"}, {"key": "executor", "label": "检修执行人"}, {"key": "reviewer", "label": "检修复核人"}, {"key": "business_participant", "label": "业务系统参与人"}, {"key": "security_officer", "label": "安全责任人"}], "rows": [{"provider": state.get("provider", ""), "executor": state.get("executor", ""), "reviewer": state.get("reviewer", ""), "business_participant": "不涉及", "security_officer": state.get("security_officer", "")}]},
                        ],
                    },
                    {
                        "heading": "风险评估",
                        "blocks": [
                            {"type": "heading", "text": "5.1影响范围", "level": 2},
                            {"type": "paragraph", "text": f"{item_name}对业务无影响；"},
                            {"type": "heading", "text": "5.2危险点分析", "level": 2},
                            {"type": "paragraphs", "items": [
                                "（1）授权不当危险点：授权过大，导致操作影响预定方案以外的生产环境实例。",
                                "（2）备份不当危险点：本次创建新实例不涉及业务数据备份，但需保留创建参数用于回滚删除核对。",
                                "（3）验证不当危险点：ECS操作对象，以及组织、资源集、VPC、VSwitch、安全组、镜像、规格、IP地址余量未核实清楚，导致创建失败或业务不可达。",
                                "（4）双人复核不当危险点：双人复核不仔细，导致操作错误执行而出现业务影响。",
                            ]},
                            {"type": "heading", "text": "5.3安全措施", "level": 2},
                            {"type": "heading", "text": "5.3.1授权", "level": 3},
                            {"type": "paragraphs", "items": [f"ASCM：国网总部直属单位权限     授权账号：{state.get('ascm_account', '')}", f"堡垒机账号：{state.get('bastion_account', '')}"]},
                            {"type": "heading", "text": "5.3.2备份", "level": 3},
                            {"type": "paragraph", "text": f"(1){item_name}不涉及备份"},
                            {"type": "heading", "text": "5.3.3验证", "level": 3},
                            {"type": "paragraph", "text": f"(1){item_name}检查资源集IP充足，确认VPC、VSwitch、安全组、镜像、规格、磁盘、数量与需求一致。"},
                            {"type": "heading", "text": "5.3.4 双人复核", "level": 3},
                            {"type": "paragraphs", "items": ["(1)确认在正确的组织和资源集下做操作，检查实例操作对象是否正确；", "(2)严格按照文档复核关键步骤及关键点。"]},
                        ],
                    },
                    {
                        "heading": "实施步骤",
                        "blocks": [
                            {"type": "heading", "text": "6.1备份", "level": 2},
                            {"type": "numbered_list", "items": [f"{item_name}不涉及备份"]},
                            {"type": "heading", "text": "6.2 检修前验证", "level": 2},
                            {"type": "numbered_list", "items": [f"{item_name}检查资源集IP充足", f"确认组织“{organization}”、资源集“{resource_set}”、云环境“{state.get('network', '')}”与工单需求一致。", "确认VPC、VSwitch、安全组、镜像、实例规格、磁盘、数量已由双人复核。"]},
                            {"type": "heading", "text": "6.3 检修操作", "level": 2},
                            {"type": "heading", "text": f"6.3.1 {item_name}", "level": 3},
                            {"type": "paragraphs", "items": [f"使用{state.get('ascm_account', '')}账号，登录{state.get('network', '')}ASCM平台，产品选择“云服务器ECS”。", f"选择组织“{organization}”-资源集“{resource_set}”，进入云服务器ECS实例列表。", "点击“创建”或“创建ECS实例”。", "按参数表填写云环境、实例名称、磁盘、自定义镜像、VPC、VSwitch、安全组、实例规格、数量等内容。", "确认填写内容无误后，由检修复核人进行关键参数复核。", "复核通过后点击提交，等待创建任务完成。", "创建完成后进入实例列表，确认新建ECS实例状态为运行中或正常，记录实例ID、IP地址和所属资源集。"]},
                            {"type": "heading", "text": "6.4 检修后验证", "level": 2},
                            {"type": "numbered_list", "items": [f"验证{item_name}实例状态正常；", "核对实例名称、规格、磁盘、镜像、VPC、VSwitch、安全组、IP、资源集与参数表一致；", "联系项目组验证业务正常；", "保留ASCM创建结果截图、实例列表截图和项目组验证记录。"]},
                        ],
                    },
                    {
                        "heading": "回滚步骤",
                        "blocks": [
                            {"type": "heading", "text": "7.1 回滚操作", "level": 2},
                            {"type": "numbered_list", "items": [f"{item_name}回退", "删除新建ecs实例", "确认实例列表中已不存在本次新建实例，或实例处于已释放状态。"]},
                            {"type": "heading", "text": "7.2 回滚后验证", "level": 2},
                            {"type": "numbered_list", "items": [f"{item_name}回退验证", "验证回退正常；联系项目组验证业务正常。"]},
                        ],
                    },
                ],
            },
        }

    return {
        "title": title,
        "department": "云运营中心平台运维处",
        "date": datetime.now().strftime("%Y年%m月%d日"),
        "evidence": {
            "selected_skills": orchestration.get("selected_skill_names", []),
            "rag_enabled": orchestration.get("rag_enabled", False),
            "rag_chunks_count": orchestration.get("rag_chunks_count", 0),
            "fallback_used": True,
            "raw_output_saved": bool(raw_text),
        },
        "document": {
            "title": title,
            "header": [
                {"text": "云运营中心平台运维处", "align": "center"},
                {"text": datetime.now().strftime("%Y年%m月%d日"), "align": "center"},
            ],
            "sections": [
                {
                    "heading": "一、背景",
                    "blocks": [
                        {
                            "type": "numbered_list",
                            "items": [state.get("background") or "根据业务运维需求，需要制定并执行本次检修方案。"],
                            "first_line_indent": 0.74,
                        },
                        {
                            "type": "paragraph",
                            "text": "以上事项由项目组提出检修需求，需通过规范化实施完成问题闭环。",
                            "first_line_indent": 0.74,
                        },
                    ],
                },
                {
                    "heading": "二、检修类型",
                    "blocks": [{"type": "checkbox_group", "items": checkbox_items, "per_line": 2}],
                },
                {
                    "heading": "三、现场环境",
                    "blocks": [
                        {
                            "type": "key_values",
                            "items": [
                                {"label": "网络环境", "value": state.get("network", "")},
                                {"label": "实施地点", "value": state.get("location", "")},
                                {"label": "检修窗口", "value": schedule_text},
                            ],
                        },
                        {"type": "paragraph", "text": "涉及实例信息：", "first_line_indent": 0.74},
                        {"type": "paragraph", "text": state.get("instances") or "待实施前由执行人员再次确认。", "first_line_indent": 0.74},
                    ],
                },
                {
                    "heading": "四、实施计划",
                    "blocks": [
                        {
                            "type": "table",
                            "columns": [
                                {"key": "role", "label": "角色"},
                                {"key": "name", "label": "人员/账号"},
                            ],
                            "rows": [
                                {"role": "方案提供人", "name": state.get("provider", "")},
                                {"role": "检修执行人", "name": state.get("executor", "")},
                                {"role": "检修复核人", "name": state.get("reviewer", "")},
                                {"role": "安全责任人", "name": state.get("security_officer", "")},
                                {"role": "ASCM授权账号", "name": state.get("ascm_account", "")},
                                {"role": "堡垒机账号", "name": state.get("bastion_account", "")},
                            ],
                        }
                    ],
                },
                {
                    "heading": "五、风险评估",
                    "blocks": [
                        {"type": "heading", "text": "5.1 危险点分析", "level": 2},
                        {"type": "plain_list", "prefix": "（1）", "items": ["检修操作可能影响相关云资源或业务访问，需要在窗口期内执行并做好监控。"]},
                        {"type": "heading", "text": "5.2 预控措施", "level": 2},
                        {"type": "plain_list", "prefix": "（1）", "items": ["实施前完成资源状态、权限、备份/快照、监控告警和回滚条件确认。"]},
                        {"type": "heading", "text": "5.3 应急处置", "level": 2},
                        {"type": "plain_list", "prefix": "（1）", "items": ["若出现异常，立即停止后续操作，保留现场信息，按回滚步骤恢复并通知相关责任人。"]},
                    ],
                },
                {
                    "heading": "六、实施步骤",
                    "blocks": [
                        {
                            "type": "numbered_list",
                            "items": [
                                "实施前确认检修对象、窗口期、授权账号和审批工单均已满足要求。",
                                f"依据检修类型“{maintenance_type}”读取对应 Skill 的实施要求，并结合 RAG 参考资料校验操作边界。",
                                f"按检修目标执行操作：{ops_hint}",
                                f"核对关键技术参数：{tech_hint}",
                                "实施完成后检查资源状态、业务连通性、监控告警和日志，确认无异常后关闭检修。",
                            ],
                            "first_line_indent": 0.74,
                        }
                    ],
                },
                {
                    "heading": "七、回滚步骤",
                    "blocks": [
                        {
                            "type": "numbered_list",
                            "items": [
                                "触发回滚条件时立即停止后续变更操作，通知复核人和安全责任人。",
                                "根据实施前确认的备份、快照、原配置或资源状态执行恢复。",
                                "回滚后重新验证业务访问、资源状态、监控告警和日志，形成处置记录。",
                            ],
                            "first_line_indent": 0.74,
                        }
                    ],
                },
            ],
        },
    }


def build_user_prompt(
    background: str, maintenance_type: str, network: str, location: str,
    instances: str, schedule_year: str, schedule_start: str, schedule_end: str,
    provider: str, executor: str, reviewer: str, security_officer: str,
    ascm_account: str, bastion_account: str, ops_detail: str, tech_params: str,
    orchestration_context: str = "",
    edit_instruction: str = "",
    previous_document_text: str = "",
) -> str:
    """Convert form fields into the user task for the Agent."""
    edit_context = ""
    if edit_instruction:
        previous_text = previous_document_text[:8000] if previous_document_text else "No previous document text was available."
        edit_context = f"""
## Document revision task
This is a revision of the previously generated maintenance plan, not a brand-new plan.
Keep the existing structure and unchanged content as much as possible. Only update the
personnel, risks, steps, rollback content, scripts, or other parts explicitly requested.

## Revision request
{edit_instruction}

## Previous document text for reference
{previous_text}
"""
    return f"""请根据以下检修需求生成标准化检修方案 JSON。

你必须先根据系统中注册的 Agent Skills 判断需要哪些 Skill，然后使用 read_file 工具读取对应 SKILL.md。若涉及多个检修类型，应组合多个 Skill。

{orchestration_context}
{edit_context}

## 背景与检修事项
{background}

## 检修类型
{maintenance_type}

## 现场环境
内/外网环境：{network}
实施地点：{location}
涉及的组件实例描述：
{instances}

## 检修窗口
{schedule_year} {schedule_start} 至 {schedule_end}

## 人员信息
方案提供人：{provider}
检修执行人：{executor}
检修复核人：{reviewer}
安全责任人：{security_officer}

## 授权账号
ASCM账号：{ascm_account}
堡垒机账号：{bastion_account}

## 检修操作补充说明（可选）
{ops_detail}

注意：即使补充说明为空，也必须根据已激活 Skill 自动生成完整的“六、实施步骤”，不要要求用户手工提供详细操作步骤。

## 技术参数
{tech_params}
"""


async def create_plan_agent() -> ReActAgent:
    """Create the native AgentScope ReActAgent for plan generation."""
    agent = ReActAgent(
        name="MaintenancePlanGenerator",
        sys_prompt=build_system_prompt(),
        model=get_model(),
        formatter=get_formatter(),
        toolkit=await get_toolkit(),
        memory=InMemoryMemory(),
        knowledge=await get_agent_knowledge(),
        enable_rewrite_query=True,
        max_iters=int(os.getenv("AGENT_MAX_ITERS", "8")),
    )
    agent.set_console_output_enabled(False)
    return agent


def format_agent_trace(msg: Msg, last: bool = True) -> list[str]:
    traces: list[str] = []
    try:
        blocks = msg.get_content_blocks()
    except Exception:
        return traces
    for block in blocks:
        if not isinstance(block, dict):
            continue
        block_type = block.get("type")
        if block_type == "tool_use":
            tool_input = json.dumps(block.get("input", {}), ensure_ascii=False)
            traces.append(f"调用工具：{block.get('name', 'unknown')}\n参数：{tool_input[:600]}")
        elif block_type == "tool_result":
            output = get_response_text(block.get("output", ""))
            traces.append(f"工具返回：{block.get('name', 'unknown')}\n{output[:800]}")
        elif block_type == "thinking":
            traces.append(f"模型思考：{str(block.get('thinking', ''))[:800]}")
        elif block_type == "text":
            text = str(block.get("text", "")).strip()
            if text:
                label = "模型输出完成" if last else "模型输出中"
                traces.append(f"{label}：{text[:800]}")
    return traces


async def run_plan_agent(user_prompt: str, trace_callback=None):
    """Run the native AgentScope ReActAgent for plan generation."""
    agent = await create_plan_agent()
    if not trace_callback:
        return await agent(Msg("user", user_prompt, "user"))

    queue: asyncio.Queue = asyncio.Queue(maxsize=100)
    agent.set_msg_queue_enabled(True, queue)
    task = asyncio.create_task(agent(Msg("user", user_prompt, "user")))

    while True:
        if task.done() and queue.empty():
            break
        try:
            msg, last, _speech = await asyncio.wait_for(queue.get(), timeout=0.2)
        except asyncio.TimeoutError:
            continue
        for trace in format_agent_trace(msg, last):
            await trace_callback(trace)
    return await task


async def run_plan_validation_agent(
    state: dict[str, str],
    filename: str,
    download_url: str,
) -> str:
    """Use Page Agent MCP to validate the generated plan from a browser."""
    agent = ReActAgent(
        name="MaintenancePlanValidator",
        sys_prompt=(
            "你是检修方案浏览器验证助手。你只能做只读验证和页面操作验证，"
            "不得执行任何真实生产变更、删除、重启、扩缩容、创建资源等不可逆操作。"
            "如果 Page Agent Hub 未连接，直接说明需要在 Chrome 中允许 Page Agent 扩展连接。"
        ),
        model=get_model(),
        formatter=get_formatter(),
        toolkit=await get_toolkit(),
        memory=InMemoryMemory(),
        max_iters=int(os.getenv("VALIDATION_AGENT_MAX_ITERS", "6")),
    )
    agent.set_console_output_enabled(False)
    prompt = f"""请通过 Page Agent MCP 对刚生成的检修方案做浏览器侧验证。

验证范围：
1. 先调用 Page Agent 状态工具确认浏览器 Hub 是否连接。
2. 如果已连接，打开 http://127.0.0.1:8000{download_url}，确认检修方案文档可访问或可下载。
3. 不要执行文档中的真实检修命令，不要登录生产系统，不要进行任何实际资源变更。
4. 输出验证结论、发现的问题和下一步人工检查建议。

生成文件：{filename}
检修类型：{state.get("maintenance_type", "")}
检修背景：{state.get("background", "")}
涉及实例：{state.get("instances", "")}
"""
    response = await agent(Msg("user", prompt, "user"))
    return get_response_text(response)


def tool_response_text(response: ToolResponse) -> str:
    parts = []
    for block in response.content:
        text = getattr(block, "text", None)
        if text is None and isinstance(block, dict):
            text = block.get("text")
        if text:
            parts.append(str(text))
    return "\n".join(parts).strip()


async def run_page_agent_task(task: str) -> str:
    task = task.strip()
    if not task:
        raise HTTPException(status_code=400, detail="请输入 Page Agent 测试指令")
    toolkit = await get_toolkit()
    tool_name = "execute_task"
    if tool_name not in toolkit.tools:
        raise HTTPException(
            status_code=503,
            detail="Page Agent MCP 工具未注册，请检查 backend/mcp_servers.json。",
        )
    tool_call = {
        "type": "tool_use",
        "id": uuid.uuid4().hex,
        "name": tool_name,
        "input": {"task": task},
    }
    result = ""
    tool_result = toolkit.call_tool_function(tool_call)
    if inspect.isawaitable(tool_result):
        tool_result = await tool_result
    async for chunk in tool_result:
        result = tool_response_text(chunk) or result
    return result or "Page Agent 执行完成，但未返回文本结果。"


# ── API 路由 ────────────────────────────────────────────────────


def detect_chat_intent(message: str, session: dict[str, Any]) -> str:
    plan_keywords = (
        "检修方案",
        "生成方案",
        "生成一个",
        "出一份",
        "写一份",
        "创建ECS",
        "创建 ECS",
        "ecs",
        "ECS",
        "rds",
        "RDS",
        "redis",
        "Redis",
        "slb",
        "SLB",
        "oss",
        "OSS",
        "k8s",
        "K8S",
        "polardb",
        "PolarDB",
    )
    if not session.get("generated"):
        return "generate" if any(keyword in message for keyword in plan_keywords) else "chat"
    regenerate_keywords = (
        "重新生成",
        "重新出",
        "再生成",
        "重做",
        "重新来",
        "重新写",
        "生成一遍",
        "再来一版",
        "按原需求",
        "根据需求重新",
        "不是修订",
        "不是修改",
    )
    if any(keyword in message for keyword in regenerate_keywords):
        return "regenerate"
    edit_keywords = (
        "修改",
        "更改",
        "调整",
        "替换",
        "修订",
        "重新评估",
        "风险点",
        "人员名单",
        "检修人员",
        "执行人",
        "复核人",
        "安全责任人",
        "按照",
        "参考",
    )
    if any(keyword in message for keyword in edit_keywords):
        return "edit"
    return "generate" if any(keyword in message for keyword in plan_keywords) else "chat"


def get_workflow_skill_body() -> str:
    skill = get_skill_registry().get("maintenance-plan-workflow")
    return skill.body if skill else ""


def should_extract_for_intent(message: str, intent: str) -> bool:
    """Regeneration-only commands should not overwrite collected requirements."""
    if intent != "regenerate":
        return True
    if "\n" in message:
        return True
    if any(mark in message for mark in ("：", ":", "；", ";")) and len(message) > 30:
        return True
    field_words = (
        "检修背景",
        "检修类型",
        "网络环境",
        "实施地点",
        "涉及实例",
        "检修窗口",
        "方案提供人",
        "检修执行人",
        "检修复核人",
        "安全责任人",
        "ASCM",
        "堡垒机",
        "技术参数",
    )
    return any(word in message for word in field_words)


async def classify_chat_intent(message: str, session: dict[str, Any]) -> dict[str, Any]:
    """Classify the latest user message before entering the plan workflow."""
    fallback_intent = detect_chat_intent(message, session)
    fallback = {
        "intent": fallback_intent,
        "should_extract": False if fallback_intent == "chat" else should_extract_for_intent(message, fallback_intent),
        "reason": "fallback",
    }
    state_preview = {
        key: value
        for key, value in session.get("state", {}).items()
        if isinstance(value, str) and value.strip()
    }
    workflow_skill = get_workflow_skill_body()
    prompt = f"""你是检修方案生成系统的意图识别器。请判断用户最新消息应该走哪个流程。

请严格遵守下面的总控工作流 Skill：
{workflow_skill[:6000]}

只能返回 JSON，不要输出 markdown，不要解释。

可选 intent：
- chat：普通交流、询问系统能力、询问原理、打招呼、咨询配置/流程/如何使用；不生成或修改文档。
- generate：用户提供检修需求，或明确要求生成新的检修方案。
- regenerate：会话中已经有生成结果，用户要求“重新生成/再生成/重做/按原需求生成一遍”，且不是要求修改上一版内容。
- edit：会话中已经有生成结果，用户要求修改上一版文档，例如变更人员、替换内容、重新评估风险点、按某文档调整已有方案。

判断规则：
- 没有已生成文档时，不要返回 edit 或 regenerate。
- “重新生成一遍/根据需求重新生成”优先是 regenerate，不是 edit。
- “修改/调整/替换/重新评估风险点/变更检修人员名单/按照某文档修订”才是 edit。
- 普通聊天不要抽取需求字段。
- regenerate 如果只是短指令，不要抽取需求字段；如果用户同时贴了新的完整需求，可以抽取。

返回格式：
{{
  "intent": "chat|generate|regenerate|edit",
  "should_extract": true,
  "reason": "一句话说明"
}}

会话是否已有生成文档：{bool(session.get("generated"))}
当前已收集字段：{json.dumps(state_preview, ensure_ascii=False)}
用户最新消息：{message}
"""
    try:
        response = await get_model()(
            [
                {"role": "system", "content": "你只输出严格 JSON。"},
                {"role": "user", "content": prompt},
            ],
        )
        data = extract_json(get_response_text(response))
    except Exception:
        return fallback

    intent = str(data.get("intent", "")).strip().lower()
    if intent not in {"chat", "generate", "regenerate", "edit"}:
        return fallback
    if not session.get("generated") and intent in {"edit", "regenerate"}:
        intent = "generate" if intent == "regenerate" else "chat"
    should_extract = data.get("should_extract")
    if not isinstance(should_extract, bool):
        should_extract = should_extract_for_intent(message, intent)
    if intent == "chat":
        should_extract = False
    return {
        "intent": intent,
        "should_extract": should_extract,
        "reason": str(data.get("reason", ""))[:200],
    }


async def run_normal_chat(session: dict[str, Any], message: str) -> str:
    """Answer normal user messages without entering the plan generation chain."""
    state_preview = {
        key: value
        for key, value in session.get("state", {}).items()
        if isinstance(value, str) and value.strip()
    }
    history = session.get("history", [])[-8:]
    history_text = "\n".join(
        f"{item.get('role', 'user')}: {item.get('content', '')}"
        for item in history
    )
    prompt = f"""你是检修方案生成系统的助手，可以正常交流，也可以说明系统如何生成、修订和验证检修方案。

当前已收集的方案字段（如有）：{json.dumps(state_preview, ensure_ascii=False)}
最近对话：
{history_text}

请直接回答用户最新问题。若用户是在询问如何生成方案，可以简要说明需要提供哪些信息；不要在普通聊天中生成 DOCX。

用户最新问题：{message}
"""
    response = await get_model()(
        [
            {"role": "system", "content": "你是检修方案生成系统的中文助手，回答要简洁准确。"},
            {"role": "user", "content": prompt},
        ],
    )
    return get_response_text(response).strip() or "我在。你可以直接描述检修需求，也可以问我这个系统怎么工作。"


def get_generated_path(session: dict[str, Any]) -> Optional[Path]:
    generated = session.get("generated") or {}
    file_id = generated.get("file_id")
    path = _generated_files.get(file_id)
    if path and path.exists():
        return path
    return None


@asynccontextmanager
async def lifespan(_app: FastAPI):
    global _rag_enabled
    registry = get_skill_registry()
    print(f"[AgentScope] 已注册 Skills: {[skill.name for skill in registry.skills]}")
    _rag_enabled = get_knowledge_base(SKILLS_ROOT) is not None
    print(f"[AgentScope] RAG enabled: {_rag_enabled}")
    print(f"[AgentScope] MCP servers configured: {[item.get('name') for item in load_mcp_server_configs()]}")
    try:
        yield
    finally:
        await close_mcp_clients()


# ── FastAPI app ─────────────────────────────────────────────────
app = FastAPI(title="检修方案生成器", version="0.4.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "skills_loaded": [skill.name for skill in get_skill_registry().skills],
        "framework": "agentscope",
        "rag_enabled": _rag_enabled,
        "mcp_servers_configured": len(load_mcp_server_configs()),
        "model_provider": os.getenv("MODEL_PROVIDER", "deepseek"),
        "model_name": os.getenv("MODEL_NAME", "deepseek-v4-pro"),
    }


@app.get("/api/mcp")
async def list_mcp_servers():
    servers = []
    for item in load_mcp_server_configs():
        servers.append(
            {
                "name": item.get("name"),
                "type": item.get("type", "http_stateless"),
                "transport": item.get("transport"),
                "url": item.get("url"),
                "command": item.get("command"),
                "args": item.get("args"),
                "cwd": item.get("cwd"),
                "group_name": item.get("group_name", "mcp"),
                "enable_funcs": item.get("enable_funcs"),
                "disable_funcs": item.get("disable_funcs"),
            }
        )
    return {"servers": servers, "count": len(servers)}


@app.post("/api/mcp/start")
async def start_mcp_servers():
    await get_toolkit()
    return {
        "status": "ok",
        "message": "MCP 已启动或已处于可用状态。",
        "servers": [item.get("name") for item in load_mcp_server_configs()],
    }


@app.post("/api/page-agent/task")
async def execute_page_agent_task(request: PageAgentTaskRequest):
    return {
        "status": "ok",
        "result": await run_page_agent_task(request.task),
    }


@app.get("/api/skills")
async def list_skills():
    registry = get_skill_registry()
    return {
        "skills": [
            {
                "name": skill.name,
                "description": skill.description,
                "path": str(skill.path),
            }
            for skill in registry.skills
        ]
    }


def safe_skill_dir_name(name: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_.-]+", "-", name.strip()).strip(".-")
    if not value:
        raise HTTPException(status_code=400, detail="Skill 名称不能为空")
    return value


def safe_relative_dir(name: str) -> str:
    value = name.strip().replace("\\", "/").strip("/")
    if not value:
        return ""
    parts = [safe_skill_dir_name(part) for part in value.split("/") if part.strip()]
    return "/".join(parts)


def safe_file_stem(name: str) -> str:
    value = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fff]+", "-", name.strip()).strip(".-")
    return value or "knowledge"


def ensure_within_directory(base: Path, target: Path) -> None:
    base_resolved = base.resolve()
    target_resolved = target.resolve()
    if not str(target_resolved).startswith(str(base_resolved)):
        raise HTTPException(status_code=400, detail="非法文件路径")


def docx_to_markdown(raw: bytes) -> str:
    from docx import Document

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(raw)
        tmp_path = Path(tmp.name)
    try:
        document = Document(str(tmp_path))
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
                lines.append("")
        for table_idx, table in enumerate(document.tables, start=1):
            lines.append(f"表格 {table_idx}")
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if rows:
                width = max(len(row) for row in rows)
                normalized = [row + [""] * (width - len(row)) for row in rows]
                lines.append("| " + " | ".join(normalized[0]) + " |")
                lines.append("| " + " | ".join(["---"] * width) + " |")
                for row in normalized[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        return "\n".join(lines).strip() + "\n"
    finally:
        tmp_path.unlink(missing_ok=True)


def docx_path_to_markdown(path: Path) -> str:
    return docx_to_markdown(path.read_bytes())


def list_knowledge_documents() -> list[dict[str, Any]]:
    knowledge_root = ROOT / "knowledge"
    if not knowledge_root.exists():
        return []
    docs = []
    for path in sorted([*knowledge_root.glob("**/*.md"), *knowledge_root.glob("**/*.txt")]):
        if path.name == ".gitkeep":
            continue
        docs.append(
            {
                "name": path.name,
                "path": str(path.relative_to(knowledge_root)),
                "size": path.stat().st_size,
                "updated_at": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            }
        )
    return docs


@app.post("/api/skills/upload")
async def upload_skill(
    file: UploadFile = File(...),
    skill_name: str = Form(default=""),
):
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件为空")

    if suffix == ".zip":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
            tmp.write(raw)
            tmp_path = Path(tmp.name)
        try:
            with zipfile.ZipFile(tmp_path) as archive:
                skill_entries = [
                    name
                    for name in archive.namelist()
                    if not name.endswith("/") and Path(name).name == "SKILL.md"
                ]
                if not skill_entries:
                    raise HTTPException(status_code=400, detail="zip 中未找到 SKILL.md")
                skill_root_parts = Path(skill_entries[0]).parent.parts
                inferred_name = skill_root_parts[-1] if skill_root_parts else Path(filename).stem
                target_dir = SKILLS_ROOT / safe_skill_dir_name(skill_name or inferred_name)
                target_dir.mkdir(parents=True, exist_ok=True)
                for member in archive.infolist():
                    if member.is_dir():
                        continue
                    member_path = Path(member.filename)
                    if ".." in member_path.parts:
                        raise HTTPException(status_code=400, detail="zip 中包含非法路径")
                    parts = member_path.parts
                    if skill_root_parts and parts[: len(skill_root_parts)] == skill_root_parts:
                        parts = parts[len(skill_root_parts) :]
                    if not parts:
                        continue
                    output_path = target_dir.joinpath(*parts)
                    ensure_within_directory(target_dir, output_path)
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                    output_path.write_bytes(archive.read(member))
        finally:
            tmp_path.unlink(missing_ok=True)
    else:
        text = raw.decode("utf-8")
        if "name:" not in text and "# " not in text:
            raise HTTPException(status_code=400, detail="请上传有效的 SKILL.md")
        target_name = safe_skill_dir_name(skill_name or Path(filename).stem or "uploaded-skill")
        target_dir = SKILLS_ROOT / target_name
        target_dir.mkdir(parents=True, exist_ok=True)
        (target_dir / "SKILL.md").write_text(text, encoding="utf-8")

    await reset_skill_runtime()
    return {
        "status": "ok",
        "skills": [
            {
                "name": skill.name,
                "description": skill.description,
                "path": str(skill.path),
            }
            for skill in get_skill_registry().skills
        ],
    }


@app.get("/api/knowledge")
async def list_knowledge():
    return {
        "documents": list_knowledge_documents(),
        "supported_extensions": [".md", ".txt", ".docx"],
        "rag_enabled": get_knowledge_base(SKILLS_ROOT) is not None,
        "chunk_size": int(os.getenv("RAG_CHUNK_SIZE", "800")),
        "split_by": os.getenv("RAG_SPLIT_BY", "paragraph"),
        "top_k": int(os.getenv("RAG_TOP_K", "5")),
        "score_threshold": os.getenv("RAG_SCORE_THRESHOLD"),
    }


@app.post("/api/knowledge/upload")
async def upload_knowledge(
    file: UploadFile = File(...),
    category: str = Form(default="uploaded"),
):
    filename = file.filename or ""
    suffix = Path(filename).suffix.lower()
    if suffix not in {".md", ".txt", ".docx"}:
        raise HTTPException(status_code=400, detail="仅支持 .md、.txt、.docx 知识文档")
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="上传文件为空")

    knowledge_root = ROOT / "knowledge"
    target_dir = knowledge_root / safe_relative_dir(category or "uploaded")
    target_dir.mkdir(parents=True, exist_ok=True)
    ensure_within_directory(knowledge_root, target_dir)

    if suffix == ".docx":
        content = docx_to_markdown(raw)
        target_path = target_dir / f"{safe_file_stem(Path(filename).stem)}.md"
    else:
        content = raw.decode("utf-8")
        target_path = target_dir / f"{safe_file_stem(Path(filename).stem)}{suffix}"
    ensure_within_directory(knowledge_root, target_path)
    target_path.write_text(content, encoding="utf-8")

    reset_knowledge_base()
    return {
        "status": "ok",
        "path": str(target_path.relative_to(knowledge_root)),
        "documents": list_knowledge_documents(),
    }


@app.post("/api/rag/reindex")
async def reindex_rag():
    reset_knowledge_base()
    knowledge_base = get_knowledge_base(SKILLS_ROOT)
    if knowledge_base is None:
        return {
            "status": "disabled",
            "message": "RAG 未启用：请配置 OPENAI_API_KEY 或 EMBEDDING_API_KEY。",
        }
    await knowledge_base.get_knowledge()
    return {
        "status": "ok",
        "documents": list_knowledge_documents(),
    }


@app.get("/api/rag/retrieve")
async def retrieve_rag(query: str = Query(default="", description="检索问题")):
    knowledge_base = get_knowledge_base(SKILLS_ROOT)
    if knowledge_base is None:
        return {
            "enabled": False,
            "chunks": [],
            "message": "RAG 未启用：请配置 OPENAI_API_KEY 或 EMBEDDING_API_KEY。",
        }
    return {
        "enabled": True,
        "chunks": await knowledge_base.retrieve(query),
    }


@app.post("/api/chat")
async def chat(request: ChatRequest):
    session_id = request.session_id or uuid.uuid4().hex
    session = _chat_sessions.setdefault(
        session_id,
        {
            "state": default_form_state(),
            "history": [],
            "generated": None,
        },
    )
    message = request.message.strip()
    if not message:
        raise HTTPException(status_code=400, detail="请输入需求描述或补充信息")

    session["history"].append({"role": "user", "content": message})
    classification = await classify_chat_intent(message, session)
    intent = classification["intent"]
    if intent == "chat":
        assistant_message = await run_normal_chat(session, message)
        session["history"].append({"role": "assistant", "content": assistant_message})
        return {
            "session_id": session_id,
            "status": "chat",
            "message": assistant_message,
            "intent": classification,
            "collected": session["state"],
        }
    extracted = {}
    if classification.get("should_extract", True):
        extracted = await extract_chat_updates(session["state"], message)
        merge_updates(session["state"], extracted.get("updates", {}))
    missing = find_missing_fields(session["state"])

    if missing and intent != "edit":
        assistant_message = extracted.get("assistant_note") or "已收到，我先整理这些信息。"
        assistant_message = f"{assistant_message}\n\n{build_missing_question(missing)}"
        session["history"].append({"role": "assistant", "content": assistant_message})
        return {
            "session_id": session_id,
            "status": "need_more",
            "message": assistant_message,
            "missing_fields": missing,
            "collected": session["state"],
        }

    orchestration = await build_generation_orchestration_context(session["state"])
    previous_document_text = ""
    if intent == "edit":
        generated_path = get_generated_path(session)
        if generated_path:
            previous_document_text = docx_path_to_markdown(generated_path)
    file_id, _path, filename = await generate_docx_from_state(
        session["state"],
        orchestration=orchestration,
        edit_instruction=message if intent == "edit" else "",
        previous_document_text=previous_document_text,
    )
    download_url = f"/api/download/{file_id}"
    assistant_message = "关键信息已收集完整，检修方案已生成。"
    if intent == "edit":
        assistant_message = "已基于上一版文档生成修订版。"
    session["generated"] = {
        "file_id": file_id,
        "filename": filename,
        "download_url": download_url,
    }
    validation_result = None
    if request.execute_validation:
        validation_result = await run_plan_validation_agent(
            session["state"],
            filename,
            download_url,
        )
    session["history"].append({"role": "assistant", "content": assistant_message})
    return {
        "session_id": session_id,
        "status": "generated",
        "message": assistant_message,
        "download_url": download_url,
        "filename": filename,
        "validation_result": validation_result,
        "collected": session["state"],
        "evidence": {
            "selected_skills": orchestration["selected_skill_names"],
            "rag_enabled": orchestration["rag_enabled"],
            "rag_chunks_count": orchestration["rag_chunks_count"],
        },
    }


def sse_event(event: str, data: dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest):
    async def stream():
        session_id = request.session_id or uuid.uuid4().hex
        session = _chat_sessions.setdefault(
            session_id,
            {
                "state": default_form_state(),
                "history": [],
                "generated": None,
            },
        )
        message = request.message.strip()
        if not message:
            yield sse_event(
                "error",
                {"session_id": session_id, "message": "请输入需求描述或补充信息"},
            )
            return

        try:
            yield sse_event(
                "status",
                {"session_id": session_id, "message": "正在识别用户意图..."},
            )
            session["history"].append({"role": "user", "content": message})
            classification = await classify_chat_intent(message, session)
            intent = classification["intent"]
            yield sse_event(
                "intent",
                {
                    "session_id": session_id,
                    "intent": intent,
                    "reason": classification.get("reason", ""),
                },
            )
            if intent == "chat":
                assistant_message = await run_normal_chat(session, message)
                session["history"].append({"role": "assistant", "content": assistant_message})
                yield sse_event(
                    "done",
                    {
                        "session_id": session_id,
                        "status": "chat",
                        "message": assistant_message,
                        "collected": session["state"],
                        "intent": classification,
                    },
                )
                return
            if intent == "edit":
                yield sse_event(
                    "status",
                    {"session_id": session_id, "message": "识别到这是对上一版文档的修订请求，正在读取已生成文档..."},
                )
            elif intent == "regenerate":
                yield sse_event(
                    "status",
                    {"session_id": session_id, "message": "识别到这是重新生成请求，将复用当前已收集需求，不读取上一版文档。"},
                )
            else:
                yield sse_event(
                    "status",
                    {"session_id": session_id, "message": "识别到检修方案生成需求，正在抽取关键信息..."},
                )
            extracted = {}
            if classification.get("should_extract", True):
                extracted = await extract_chat_updates(session["state"], message)
                merge_updates(session["state"], extracted.get("updates", {}))

            yield sse_event(
                "collected",
                {
                    "session_id": session_id,
                    "message": "关键信息抽取完成，正在检查是否需要补充。",
                    "collected": session["state"],
                },
            )
            missing = find_missing_fields(session["state"])

            if missing and intent != "edit":
                assistant_message = extracted.get("assistant_note") or "已收到，我先整理这些信息。"
                assistant_message = f"{assistant_message}\n\n{build_missing_question(missing)}"
                session["history"].append({"role": "assistant", "content": assistant_message})
                yield sse_event(
                    "done",
                    {
                        "session_id": session_id,
                        "status": "need_more",
                        "message": assistant_message,
                        "missing_fields": missing,
                        "collected": session["state"],
                    },
                )
                return

            yield sse_event(
                "status",
                {"session_id": session_id, "message": "信息已完整，正在初筛 Skill 并检索 RAG 参考资料..."},
            )
            orchestration = await build_generation_orchestration_context(session["state"])
            yield sse_event(
                "evidence",
                {
                    "session_id": session_id,
                    "message": "生成依据已准备完成，正在调用 AgentScope Skill 生成检修方案。",
                    "selected_skills": orchestration["selected_skill_names"],
                    "rag_enabled": orchestration["rag_enabled"],
                    "rag_chunks_count": orchestration["rag_chunks_count"],
                },
            )
            previous_document_text = ""
            if intent == "edit":
                generated_path = get_generated_path(session)
                if generated_path:
                    previous_document_text = docx_path_to_markdown(generated_path)

            trace_queue: asyncio.Queue = asyncio.Queue()

            async def trace_callback(trace_message: str) -> None:
                await trace_queue.put({"session_id": session_id, "message": trace_message})

            generation_task = asyncio.create_task(
                generate_docx_from_state(
                    session["state"],
                    orchestration=orchestration,
                    trace_callback=trace_callback,
                    edit_instruction=message if intent == "edit" else "",
                    previous_document_text=previous_document_text,
                )
            )
            while True:
                if generation_task.done() and trace_queue.empty():
                    break
                try:
                    trace_data = await asyncio.wait_for(trace_queue.get(), timeout=0.2)
                except asyncio.TimeoutError:
                    continue
                yield sse_event("trace", trace_data)

            file_id, _path, filename = await generation_task
            download_url = f"/api/download/{file_id}"
            assistant_message_override = "已基于上一版文档生成修订版。" if intent == "edit" else None
            assistant_message = "关键信息已收集完整，检修方案已生成。"
            if assistant_message_override:
                assistant_message = assistant_message_override
            session["generated"] = {
                "file_id": file_id,
                "filename": filename,
                "download_url": download_url,
            }
            session["history"].append({"role": "assistant", "content": assistant_message})
            validation_result = None
            if request.execute_validation:
                yield sse_event(
                    "status",
                    {
                        "session_id": session_id,
                        "message": "已生成检修方案，正在通过 Page Agent MCP 执行浏览器侧验证...",
                    },
                )
                validation_result = await run_plan_validation_agent(
                    session["state"],
                    filename,
                    download_url,
                )
            yield sse_event(
                "done",
                {
                    "session_id": session_id,
                    "status": "generated",
                    "message": assistant_message,
                    "download_url": download_url,
                    "filename": filename,
                    "validation_result": validation_result,
                    "collected": session["state"],
                    "evidence": {
                        "selected_skills": orchestration["selected_skill_names"],
                        "rag_enabled": orchestration["rag_enabled"],
                        "rag_chunks_count": orchestration["rag_chunks_count"],
                    },
                },
            )
        except HTTPException as exc:
            yield sse_event(
                "error",
                {"session_id": session_id, "message": exc.detail, "status_code": exc.status_code},
            )
        except Exception as exc:
            yield sse_event(
                "error",
                {"session_id": session_id, "message": f"生成失败：{exc}"},
            )

    return StreamingResponse(
        stream(),
        media_type="text/event-stream; charset=utf-8",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/api/chat/reset")
async def reset_chat(request: ChatResetRequest):
    _chat_sessions.pop(request.session_id, None)
    return {"status": "ok"}


@app.get("/api/download/{file_id}")
async def download_generated(file_id: str):
    path = _generated_files.get(file_id)
    if path is None or not path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已过期")
    return FileResponse(
        path=str(path),
        filename=path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def get_response_text(response) -> str:
    """Extract text from AgentScope 1.x ChatResponse."""
    def collect(value: Any, depth: int = 0) -> list[str]:
        if value is None or depth > 8:
            return []
        if isinstance(value, str):
            return [value]
        if isinstance(value, (int, float, bool)):
            return []
        if isinstance(value, dict):
            fragments: list[str] = []
            priority_keys = (
                "text",
                "content",
                "arguments",
                "input",
                "output",
                "message",
                "value",
                "tool_calls",
                "choices",
            )
            for key in priority_keys:
                if key in value:
                    fragments.extend(collect(value[key], depth + 1))
            for key, item in value.items():
                if key not in priority_keys:
                    fragments.extend(collect(item, depth + 1))
            if "document" in value:
                fragments.insert(0, json.dumps(value, ensure_ascii=False))
            return fragments
        if isinstance(value, (list, tuple)):
            fragments = []
            for item in value:
                fragments.extend(collect(item, depth + 1))
            return fragments

        fragments = []
        for method in ("get_text_content", "model_dump", "dict"):
            try:
                func = getattr(value, method)
            except (AttributeError, KeyError):
                continue
            if callable(func):
                try:
                    result = func()
                except TypeError:
                    continue
                fragments.extend(collect(result, depth + 1))

        for attr in ("text", "content", "message", "output", "tool_calls", "metadata"):
            try:
                attr_value = getattr(value, attr)
            except (AttributeError, KeyError):
                continue
            fragments.extend(collect(attr_value, depth + 1))
        return fragments

    seen = set()
    ordered = []
    for fragment in collect(response):
        cleaned = fragment.strip()
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            ordered.append(cleaned)
    return "\n".join(ordered)


def default_form_state() -> dict[str, str]:
    return {
        "background": "",
        "maintenance_type": "",
        "network": "",
        "location": "国网亦庄数据中心二期运维专区",
        "instances": "",
        "schedule_year": str(datetime.now().year) + "年",
        "schedule_start": "",
        "schedule_end": "",
        "provider": "",
        "executor": "",
        "reviewer": "",
        "security_officer": "",
        "ascm_account": "",
        "bastion_account": "",
        "ops_detail": "",
        "tech_params": "",
    }


def normalize_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value).strip()


def merge_updates(state: dict[str, str], updates: dict[str, Any]) -> None:
    for key in FORM_FIELDS:
        value = normalize_value(updates.get(key))
        if value:
            state[key] = value


def find_missing_fields(state: dict[str, str]) -> list[str]:
    return [key for key in REQUIRED_FIELDS if not state.get(key, "").strip()]


def build_missing_question(missing: list[str]) -> str:
    labels = [REQUIRED_FIELDS[key] for key in missing[:4]]
    if not labels:
        return ""
    return "还需要补充：" + "、".join(labels) + "。请直接回复这些信息即可。"


def infer_updates_from_text(user_message: str) -> dict[str, str]:
    """Capture obvious requirement clues before relying on free-form LLM output."""
    text = user_message.strip()
    lower_text = text.lower()
    updates: dict[str, str] = {}

    if text:
        updates["background"] = text

    def labeled_value(*labels: str, multiline: bool = False) -> str:
        label_group = "|".join(re.escape(label) for label in labels)
        if multiline:
            next_labels = (
                "检修背景|检修类型|网络环境|内外网环境|实施地点|涉及实例|检修窗口|"
                "方案提供人|检修执行人|检修复核人|安全责任人|ASCM 授权账号|"
                "ASCM授权账号|堡垒机账号|技术参数|补充要求"
            )
            pattern = rf"(?:{label_group})\s*[:：]\s*(.*?)(?=\n\s*(?:{next_labels})\s*[:：]|\Z)"
            match = re.search(pattern, text, re.S | re.I)
        else:
            match = re.search(rf"(?:{label_group})\s*[:：]\s*([^\n]+)", text, re.I)
        return match.group(1).strip() if match else ""

    labeled_background = labeled_value("检修背景", multiline=True)
    if labeled_background:
        updates["background"] = labeled_background

    label_map = {
        "maintenance_type": ("检修类型",),
        "network": ("网络环境", "内外网环境"),
        "location": ("实施地点",),
        "instances": ("涉及实例", "涉及的组件实例"),
        "provider": ("方案提供人",),
        "executor": ("检修执行人",),
        "reviewer": ("检修复核人",),
        "security_officer": ("安全责任人",),
        "ascm_account": ("ASCM 授权账号", "ASCM授权账号"),
        "bastion_account": ("堡垒机账号",),
    }
    for field, labels in label_map.items():
        value = labeled_value(*labels)
        if value:
            updates[field] = value

    tech_params = labeled_value("技术参数", multiline=True)
    if tech_params:
        updates["tech_params"] = tech_params
    ops_detail = labeled_value("补充要求", multiline=True)
    if ops_detail:
        updates["ops_detail"] = ops_detail

    schedule = labeled_value("检修窗口")
    if schedule:
        year_match = re.search(r"(\d{4}年)", schedule)
        if year_match:
            updates["schedule_year"] = year_match.group(1)
        parts = re.split(r"\s*(?:至|到|-|—|~)\s*", schedule, maxsplit=1)
        if len(parts) == 2:
            updates["schedule_start"] = parts[0].strip()
            updates["schedule_end"] = parts[1].strip()
        else:
            updates["schedule_start"] = schedule

    has_internal = "内网" in text
    has_external = "外网" in text
    if has_internal and has_external:
        updates["network"] = "内、外网"
    elif has_internal:
        updates["network"] = "内网"
    elif has_external:
        updates["network"] = "外网"

    if any(keyword in lower_text for keyword in ("ecs", "云服务器")) and any(
        keyword in text for keyword in ("创建", "新建", "申请", "开通")
    ):
        updates["maintenance_type"] = "配置变更"
        if not updates.get("instances"):
            updates["instances"] = text
    elif any(keyword in text for keyword in ("扩容", "缩容", "扩缩容")):
        updates["maintenance_type"] = "组件扩缩容"
    elif any(keyword in text for keyword in ("升级", "版本")):
        updates["maintenance_type"] = "组件升级"
    elif any(keyword in lower_text for keyword in ("数据库", "polardb", "mysql", "mongodb", "redis")):
        updates["maintenance_type"] = "数据库变更"

    return updates


async def extract_chat_updates(state: dict[str, str], user_message: str) -> dict[str, Any]:
    """Use the current LLM to update the structured requirement state."""
    prompt = f"""你是检修方案需求信息抽取助手。请从用户最新消息中抽取检修方案生成所需字段，并结合已有状态更新。

已有状态 JSON：
{json.dumps(state, ensure_ascii=False, indent=2)}

用户最新消息：
{user_message}

字段说明：
- background: 检修背景/检修事项，可多行
- maintenance_type: 配置变更/组件升级/组件扩缩容/数据库变更/日常维护（原硬件设备）/其他
- network: 内网/外网/内、外网
- location: 实施地点
- instances: 涉及实例，尽量包含事项名称、组织、资源集
- schedule_year/schedule_start/schedule_end: 检修窗口
- provider/executor/reviewer/security_officer: 人员信息
- ascm_account/bastion_account: 授权账号
- ops_detail: 用户额外约束或补充说明，不要把它当成详细步骤来源
- tech_params: 技术参数 JSON 或自然语言参数

只输出 JSON：
{{
  "updates": {{"field": "value"}},
  "assistant_note": "对已收到信息的简短确认，不超过60字"
}}

不要输出 markdown，不要解释。"""
    inferred_updates = infer_updates_from_text(user_message)
    try:
        response = await get_model()(
            [
                {"role": "system", "content": "你只输出可解析 JSON。"},
                {"role": "user", "content": prompt},
            ],
        )
    except Exception:
        return {
            "updates": inferred_updates,
            "assistant_note": "已收到需求描述，我先整理出可识别的信息。",
        }
    try:
        data = extract_json(get_response_text(response))
    except json.JSONDecodeError:
        data = {"updates": {}, "assistant_note": "已收到需求描述，我先整理出可识别的信息。"}
    if not isinstance(data, dict):
        data = {"updates": {}}

    model_updates = data.get("updates") if isinstance(data.get("updates"), dict) else {}
    data["updates"] = {**inferred_updates, **model_updates}
    if not data.get("assistant_note") and data["updates"]:
        data["assistant_note"] = "已收到需求描述，我先整理出可识别的信息。"
    return data


async def generate_docx_from_state(
    state: dict[str, str],
    orchestration: Optional[dict[str, Any]] = None,
    trace_callback=None,
    edit_instruction: str = "",
    previous_document_text: str = "",
) -> tuple[str, Path, str]:
    if orchestration is None:
        orchestration = await build_generation_orchestration_context(state)
    user_prompt = build_user_prompt(
        **state,
        orchestration_context=orchestration["prompt_context"],
        edit_instruction=edit_instruction,
        previous_document_text=previous_document_text,
    )
    response = await run_plan_agent(user_prompt, trace_callback=trace_callback)
    text = get_response_text(response)
    try:
        data = await repair_and_extract_json(text)
    except HTTPException:
        write_model_output_debug(text)
        data = build_fallback_plan_data(state, orchestration, text)

    data.setdefault("department", "云运营中心平台运维处")
    data.setdefault("date", datetime.now().strftime("%Y年%m月%d日"))
    data.setdefault("evidence", {})
    if isinstance(data["evidence"], dict):
        data["evidence"].setdefault(
            "selected_skills",
            orchestration["selected_skill_names"],
        )
        data["evidence"].setdefault("rag_enabled", orchestration["rag_enabled"])
        data["evidence"].setdefault(
            "rag_chunks_count",
            orchestration["rag_chunks_count"],
        )

    doc = build_document(data)
    output_dir = Path(tempfile.gettempdir()) / "plan-generator"
    output_dir.mkdir(parents=True, exist_ok=True)
    file_id = uuid.uuid4().hex
    output_path = output_dir / f"检修方案_{file_id[:8]}.docx"
    doc.save(str(output_path))
    _generated_files[file_id] = output_path

    files = sorted(
        output_dir.glob("检修方案_*.docx"),
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    for old_file in files[10:]:
        try:
            old_file.unlink()
        except OSError:
            pass

    filename = data.get("title") or data.get("document", {}).get("title", "检修方案")
    return file_id, output_path, filename + ".docx"


@app.post("/api/generate")
async def generate_plan(
    background: str = Form(default=""),
    maintenance_type: str = Form(default="配置变更"),
    network: str = Form(default="内网"),
    location: str = Form(default="国网亦庄数据中心二期运维专区"),
    instances: str = Form(default=""),
    schedule_year: str = Form(default=str(datetime.now().year) + "年"),
    schedule_start: str = Form(default=""),
    schedule_end: str = Form(default=""),
    provider: str = Form(default=""),
    executor: str = Form(default=""),
    reviewer: str = Form(default=""),
    security_officer: str = Form(default=""),
    ascm_account: str = Form(default=""),
    bastion_account: str = Form(default=""),
    ops_detail: str = Form(default=""),
    tech_params: str = Form(default=""),
):
    """接收检修需求，调用 Claude API 生成结构化数据，返回 .docx 文件。"""
    try:
        state = {
            "background": background,
            "maintenance_type": maintenance_type,
            "network": network,
            "location": location,
            "instances": instances,
            "schedule_year": schedule_year,
            "schedule_start": schedule_start,
            "schedule_end": schedule_end,
            "provider": provider,
            "executor": executor,
            "reviewer": reviewer,
            "security_officer": security_officer,
            "ascm_account": ascm_account,
            "bastion_account": bastion_account,
            "ops_detail": ops_detail,
            "tech_params": tech_params,
        }
        orchestration = await build_generation_orchestration_context(state)
        user_prompt = build_user_prompt(
            background=background, maintenance_type=maintenance_type,
            network=network, location=location, instances=instances,
            schedule_year=schedule_year, schedule_start=schedule_start,
            schedule_end=schedule_end, provider=provider, executor=executor,
            reviewer=reviewer, security_officer=security_officer,
            ascm_account=ascm_account, bastion_account=bastion_account,
            ops_detail=ops_detail, tech_params=tech_params,
            orchestration_context=orchestration["prompt_context"],
        )

        response = await run_plan_agent(user_prompt)

        text = get_response_text(response)
        try:
            data = await repair_and_extract_json(text)
        except HTTPException:
            write_model_output_debug(text)
            data = build_fallback_plan_data(state, orchestration, text)

        data.setdefault("department", "云运营中心平台运维处")
        data.setdefault("date", datetime.now().strftime("%Y年%m月%d日"))
        data.setdefault("evidence", {})
        if isinstance(data["evidence"], dict):
            data["evidence"].setdefault(
                "selected_skills",
                orchestration["selected_skill_names"],
            )
            data["evidence"].setdefault("rag_enabled", orchestration["rag_enabled"])
            data["evidence"].setdefault(
                "rag_chunks_count",
                orchestration["rag_chunks_count"],
            )

        doc = build_document(data)
        output_dir = Path(tempfile.gettempdir()) / "plan-generator"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"检修方案_{uuid.uuid4().hex[:8]}.docx"
        doc.save(str(output_path))

        # 清理旧文件
        files = sorted(
            output_dir.glob("检修方案_*.docx"),
            key=lambda f: f.stat().st_mtime, reverse=True,
        )
        for old_file in files[10:]:
            try:
                old_file.unlink()
            except OSError:
                pass

        filename = data.get("title") or data.get("document", {}).get("title", "检修方案")
        filename += ".docx"
        return FileResponse(
            path=str(output_path),
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="LLM 返回的数据格式异常，请重试")
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成失败: {str(e)}")


# ── 前端静态文件 ─────────────────────────────────────────────────
frontend_dir = ROOT.parent / "frontend"
if frontend_dir.exists():
    app.mount("/", StaticFiles(directory=str(frontend_dir), html=True), name="static")
