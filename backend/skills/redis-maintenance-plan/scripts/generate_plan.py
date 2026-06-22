#!/usr/bin/env python3
"""生成 Redis 实例检修方案 .docx 文件"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
# 排版常量 — 按需修改
# ============================================================

# 字体名称
FONT_FZXBS = "方正小标宋_GBK"      # 封面标题
FONT_HEITI = "黑体"                # 一级标题
FONT_KAITI = "楷体_GB2312"         # 二级标题
FONT_FANGSONG = "仿宋_GB2312"      # 正文 / 三级标题 / 封面信息
FONT_SONGTI = "宋体"               # 表格 / 页码

# 字号 (中文磅值)
SIZE_ERHAO = Pt(22)    # 二号 — 封面标题
SIZE_SANHAO = Pt(16)   # 三号 — 正文 / 各级标题 / 封面信息
SIZE_SIHAO = Pt(14)    # 四号 — 表格 / 页码
SIZE_11PT = Pt(11)     # 11pt — 现场环境实例表格

# 行距与段距
LINE_SPACING = Pt(28)                        # 固定 28 磅
PARA_SPACING_BEFORE = Pt(14)                 # 段前 0.5 行
PARA_SPACING_AFTER = Pt(14)                  # 段后 0.5 行

# 页边距
MARGIN_TOP = Cm(3.4)
MARGIN_BOTTOM = Cm(3.4)
MARGIN_LEFT = Cm(2.8)
MARGIN_RIGHT = Cm(2.8)

# Logo 设置
LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "logo.png"
LOGO_WIDTH = Cm(3.5)       # logo 宽度
LOGO_LEFT = MARGIN_LEFT    # 距页面左侧（与左边距对齐）
LOGO_TOP = MARGIN_TOP      # 距页面顶部（与上边距对齐）

# 首行缩进 (2 个三号汉字 ≈ 32pt)
FIRST_LINE_INDENT = Pt(32)

# ============================================================
# 基础排版函数
# ============================================================

def _set_font(run, font_name, font_size, bold=False, color=RGBColor(0, 0, 0)):
    """设置 run 的西文 + 东亚字体、字号、加粗、颜色"""
    run.font.size = font_size
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.bold = bold
    run.font.color.rgb = color


def _set_para_spacing(para):
    """设置段落行距（固定 28 磅）+ 段前段后"""
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = PARA_SPACING_BEFORE
    pf.space_after = PARA_SPACING_AFTER


def _new_para(doc, alignment=None, first_line_indent=None):
    """创建段落并应用默认间距"""
    para = doc.add_paragraph()
    _set_para_spacing(para)
    if alignment is not None:
        para.alignment = alignment
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = first_line_indent
    return para


def _add_run(para, text, font_name, font_size, bold=False):
    """向段落添加一个格式化 run"""
    run = para.add_run(text)
    _set_font(run, font_name, font_size, bold)
    return run


# --- 封面 ---

def add_cover_title(doc, text):
    """封面标题：方正小标宋_GBK 二号 加粗 居中"""
    para = _new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(para, text, FONT_FZXBS, SIZE_ERHAO, bold=True)
    return para


def add_cover_info(doc, text):
    """封面信息：仿宋_GB2312 三号 居中"""
    para = _new_para(doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO)
    return para


def add_cover_blank(doc):
    """封面空行（同样应用段距以保持一致）"""
    para = _new_para(doc)
    return para


def add_cover_logo(doc, logo_path, width=LOGO_WIDTH, left=LOGO_LEFT, top=LOGO_TOP):
    """在封面左上角添加公司logo（浮动定位，浮于文字上方，不占段落空间）"""
    if not logo_path.exists():
        return

    # 创建零高度段落作为图片锚点
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = para.add_run()
    run.font.size = Pt(1)
    inline_shape = run.add_picture(str(logo_path), width=width)

    # 获取 inline 元素及其父元素 w:drawing
    inline = inline_shape._inline
    drawing = inline.getparent()

    # 提取 extent
    extent_el = inline.find(qn('wp:extent'))
    cx = extent_el.get('cx')
    cy = extent_el.get('cy')

    # 提取 docPr
    docPr_el = inline.find(qn('wp:docPr'))
    docPr_id = docPr_el.get('id')
    docPr_name = docPr_el.get('name')

    # 提取 graphic（包含 blip 引用等完整图片数据）
    graphic_el = inline.find(qn('a:graphic'))

    # 移除 inline 元素
    drawing.remove(inline)

    # 构建 wp:anchor（绝对定位浮动图片）
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '0')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    # simplePos（必须字段）
    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)

    # positionH — 水平位置，相对于页面
    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'page')
    posOffsetH = OxmlElement('wp:posOffset')
    posOffsetH.text = str(int(left))
    posH.append(posOffsetH)
    anchor.append(posH)

    # positionV — 垂直位置，相对于页面
    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'page')
    posOffsetV = OxmlElement('wp:posOffset')
    posOffsetV.text = str(int(top))
    posV.append(posOffsetV)
    anchor.append(posV)

    # extent（复用 inline 中的尺寸）
    extent = OxmlElement('wp:extent')
    extent.set('cx', cx)
    extent.set('cy', cy)
    anchor.append(extent)

    # effectExtent
    effectExtent = OxmlElement('wp:effectExtent')
    effectExtent.set('l', '0')
    effectExtent.set('t', '0')
    effectExtent.set('r', '0')
    effectExtent.set('b', '0')
    anchor.append(effectExtent)

    # wrapNone — 浮于文字上方
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)

    # docPr（复用）
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', docPr_id)
    docPr.set('name', docPr_name)
    anchor.append(docPr)

    # cNvGraphicFramePr
    cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
    anchor.append(cNvGraphicFramePr)

    # graphic（复用完整图形数据）
    anchor.append(graphic_el)

    # 将 anchor 添加到 drawing 中
    drawing.append(anchor)


# --- 标题 ---

def _set_outline_level(para, level):
    """设置段落大纲级别，使其在 Word 导航窗格中显示。
    level: 0=一级, 1=二级, 2=三级"""
    pPr = para._element.get_or_add_pPr()
    outline_lvl = OxmlElement('w:outlineLvl')
    outline_lvl.set(qn('w:val'), str(level))
    pPr.append(outline_lvl)


def add_heading_1(doc, text):
    """一级标题：黑体 三号 加粗，左对齐，无缩进"""
    para = _new_para(doc, first_line_indent=Pt(0))
    _set_outline_level(para, 0)
    _add_run(para, text, FONT_HEITI, SIZE_SANHAO, bold=True)
    return para


def add_heading_2(doc, text):
    """二级标题：楷体_GB2312 三号 加粗，缩进1字符"""
    para = _new_para(doc, first_line_indent=Pt(16))
    _set_outline_level(para, 1)
    _add_run(para, text, FONT_KAITI, SIZE_SANHAO, bold=True)
    return para


def add_heading_3(doc, text):
    """三级标题：仿宋_GB2312 三号 加粗，缩进2字符"""
    para = _new_para(doc, first_line_indent=FIRST_LINE_INDENT)
    _set_outline_level(para, 2)
    _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO, bold=True)
    return para


# --- 正文 ---

def add_body(doc, text):
    """正文：仿宋_GB2312 三号，首行缩进 2 汉字"""
    para = _new_para(doc, first_line_indent=FIRST_LINE_INDENT)
    _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO)
    return para


def add_body_no_indent(doc, text):
    """正文无缩进版（用于勾选行等特殊场景）"""
    para = _new_para(doc, first_line_indent=Pt(0))
    _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO)
    return para


# --- 带格式的正文 run（用于一行内混排 ☑ / 文本） ---

def add_body_mixed(doc, segments):
    """
    segments: [(text, bold), ...]
    每个 segment 用仿宋_GB2312 三号渲染，可单独控制加粗
    """
    para = _new_para(doc, first_line_indent=FIRST_LINE_INDENT)
    for text, bold in segments:
        _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO, bold)
    return para


def add_body_mixed_no_indent(doc, segments):
    """无缩进版"""
    para = _new_para(doc, first_line_indent=Pt(0))
    for text, bold in segments:
        _add_run(para, text, FONT_FANGSONG, SIZE_SANHAO, bold)
    return para


# ============================================================
# 表格
# ============================================================

def _set_cell_paragraph(cell, text, font_name, font_size, bold=False, alignment=None):
    """清空单元格并写入格式化文本"""
    # 清空所有现有段落
    for p in cell.paragraphs:
        p.clear()
    # 使用第一个段落
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if alignment is not None:
        para.alignment = alignment
    run = para.add_run(text)
    _set_font(run, font_name, font_size, bold)


def add_table(doc, headers, rows, font_size=None):
    """
    表格：宋体，固定行距 28 磅
    表头加粗居中，数据行左对齐
    """
    if font_size is None:
        font_size = SIZE_SIHAO
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")

    # 表头
    for i, header in enumerate(headers):
        _set_cell_paragraph(table.rows[0].cells[i], header,
                           FONT_SONGTI, font_size, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            _set_cell_paragraph(table.rows[r + 1].cells[c], str(val),
                               FONT_SONGTI, font_size)

    # 表后空行（正文格式）
    add_body_no_indent(doc, "")
    return table


# ============================================================
# 页脚：页码
# ============================================================

def add_page_number(section, font_name=FONT_SONGTI, font_size=SIZE_SIHAO):
    """在 section 页脚居中添加 "—1—" 格式页码"""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # PAGE 字段
    run1 = para.add_run("—")
    _set_font(run1, font_name, font_size)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run2 = para.add_run()
    run2._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run3 = para.add_run()
    run3._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run4 = para.add_run()
    run4._r.append(fldChar2)

    run5 = para.add_run("—")
    _set_font(run5, font_name, font_size)


# ============================================================
# 实例表格规格
# ============================================================

def get_instance_table_spec(op_type):
    """根据操作类型返回实例表格规格"""
    specs = {
        "创建": {
            "title": "涉及创建的Redis实例：",
            "headers": ["数据库名称", "数据库版本", "实例规格", "芯片架构",
                       "VPC ID或名称", "Vswitch ID或名称", "密码", "数量", "环境"],
        },
        "回收": {
            "title": "涉及回收的Redis实例：",
            "headers": ["实例ID", "实例名称", "实例规格", "环境"],
        },
        "升配": {
            "title": "涉及的Redis实例：",
            "headers": ["实例名称", "实例ID", "当前规格", "目标规格", "环境"],
        },
        "降配": {
            "title": "涉及的Redis实例：",
            "headers": ["实例名称", "实例ID", "当前规格", "目标规格", "环境"],
        },
    }
    return specs.get(op_type, specs["创建"])


# ============================================================
# 默认内容生成（按类型分派）
# ============================================================

def get_default_backup(title, op_type):
    no_backup = ["创建", "回收"]
    if op_type in no_backup:
        return f"{title}不涉及备份"
    return f"{title}检查备份正常"


def get_default_verification(title, op_type):
    if op_type == "创建":
        return f"{title}检查资源集IP充足"
    return f"{title}检查实例运行正常"


# --- Section 3 & 6.3 实例信息展示 ---

def _build_instance_info(doc, op_type, instances, table_spec):
    """按操作类型展示实例信息（第3节和6.3.1下复用）"""
    if not instances:
        return
    add_body(doc, table_spec["title"])
    if op_type == "回收":
        _instances_as_text(doc, instances)
    else:
        add_table(doc, table_spec["headers"], instances, font_size=SIZE_11PT)


def _instances_as_text(doc, instances):
    """回收：纯文本行，多实例按环境分组"""
    # 按环境分组
    groups = {}
    for inst in instances:
        env = inst[3] if len(inst) > 3 else ""
        groups.setdefault(env, []).append(inst)

    for env, group in groups.items():
        if len(groups) > 1:
            add_body(doc, env)
        for inst in group:
            line = f"{inst[0]}  {inst[1]} {inst[2]}"
            add_body(doc, line)


# --- Section 6.3 检修操作 ---

def _get_op_heading(op_type, title):
    """6.3.1 小标题"""
    headings = {
        "创建": "6.3.1 创建redis实例",
        "回收": "6.3.1 回收redis实例",
    }
    return headings.get(op_type, f"6.3.1 {title}")


def _build_default_op_steps(doc, op_type, org, resource_set, account, instances):
    """按操作类型生成默认检修操作步骤"""
    if op_type == "创建":
        _op_create(doc, org, resource_set, account)
    elif op_type == "回收":
        _op_recycle(doc, org, resource_set, account, instances)
    elif op_type == "升配":
        _op_upgrade(doc, org, resource_set, account)
    elif op_type == "降配":
        _op_downgrade(doc, org, resource_set, account)


def _op_create(doc, org, resource_set, account):
    """创建：3段式（模板原文）"""
    add_body(doc,
        f'使用账号{account}登陆外网云ASCM平台，产品选择"云数据库Redis"，'
        f'选择组织"{org}"-资源集"{resource_set}"，点击"创建实例"。')
    add_body(doc, "按需选择芯片、数据库版本、架构类型、实例规格")
    add_body(doc, "网络类型选择专有网络，选择vpc和vswitch，输入实例名称、密码、创建数量，点击提交")


def _op_recycle(doc, org, resource_set, account, instances):
    """回收：2段式（模板原文）"""
    if instances:
        inst = instances[0]
        inst_ref = f"{inst[0]} {inst[1]}"
    else:
        inst_ref = "目标实例"

    add_body(doc,
        f'使用账号{account}登录ascm平台，进入云数据库Redis，'
        f'组织选择"{org}"，资源集选择"{resource_set}"，进入rds实例：{inst_ref}')
    add_body(doc, "点击释放实例，确认释放实例。等待任务完成。")


def _op_upgrade(doc, org, resource_set, account):
    """升配：3段式（模板原文）"""
    add_body(doc,
        f'使用账号{account}登陆内网云ASCM平台，产品选择"云数据库Redis"，'
        f'选择组织"{org}"-资源集"{resource_set}"，点击进入指定Redis实例-实例信息，点击"变更配置"')
    add_body(doc, '按需选择架构类型和实例规格，点击提交')
    add_body(doc, '等待任务完成。')


def _op_downgrade(doc, org, resource_set, account):
    """降配：2段式（模板原文）"""
    add_body(doc,
        f'使用账号{account}登陆内网云ASCM平台，产品选择"云数据库Redis"，'
        f'选择组织"{org}"-资源集"{resource_set}"，点击进入指定Redis实例-实例信息，点击"变更配置"')
    add_body(doc, '选择较低的目标规格，点击提交')


# --- Section 7.1 回滚 ---

def _build_default_rollback(doc, op_type, title):
    """按操作类型生成默认回滚操作"""
    add_body(doc, f"1、{title}回退")
    rb_map = {
        "创建": "删除新建redis实例",
        "回收": "回收实例不涉及回退",
    }
    add_body(doc, rb_map.get(op_type, "按照检修回退原实例"))


# ============================================================
# 主生成函数
# ============================================================

def generate_plan(data, output_path):
    doc = Document()

    # --- 页面设置 ---
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        add_page_number(section)

    title = data.get("title", "Redis实例检修方案")
    date = data.get("date", datetime.now().strftime("%Y年%m月%d日"))
    op_type = data.get("operation_type", "创建")

    # ================================================================
    # 封面
    # ================================================================
    add_cover_logo(doc, LOGO_PATH)

    for _ in range(4):
        add_cover_blank(doc)

    add_cover_title(doc, f"{title}检修方案")

    for _ in range(5):
        add_cover_blank(doc)

    add_cover_info(doc, "云运营中心平台运维处")
    add_cover_info(doc, date)

    add_cover_blank(doc)  # 过渡
    doc.add_page_break()

    # ================================================================
    # 1. 背景
    # ================================================================
    add_heading_1(doc, "1. 背景")
    bg = data.get("background", title)
    add_body(doc, bg)
    add_body(doc, "该事项项目组提报问题工单，需检修进行处理，实现问题工单闭环。")

    # ================================================================
    # 2. 检修类型
    # ================================================================
    add_heading_1(doc, "2. 检修类型")
    add_body(doc, "（勾选对应检修类型）")

    all_types_line1 = ["配置变更", "组件升级", "组件扩缩容"]
    all_types_line2 = ["数据库变更", "日常维护（原硬件设备）"]
    checked = data.get("maintenance_types", ["配置变更", "数据库变更"])

    # 第一行：配置变更、组件升级、组件扩缩容
    seg1 = []
    for mt in all_types_line1:
        mark = "☑" if mt in checked else "☐"
        seg1.append((f"  {mark} {mt}   ", False))
    add_body_mixed(doc, seg1)

    # 第二行：数据库变更、日常维护（原硬件设备）
    seg2 = []
    for mt in all_types_line2:
        mark = "☑" if mt in checked else "☐"
        seg2.append((f"  {mark} {mt}   ", False))
    add_body_mixed(doc, seg2)

    # 第三行：其他 ____
    mark_other = "☑" if "其他" in checked else "☐"
    add_body_mixed(doc, [(f"  {mark_other} 其他 ____", False)])

    # ================================================================
    # 3. 现场环境
    # ================================================================
    add_heading_1(doc, "3. 现场环境")
    env = data.get("environment", {})
    network = env.get("network", "内网")
    location = env.get("location", "国网亦庄数据中心二期运维专区")
    version = env.get("cloud_version", "v3.16")

    add_body(doc, f"（1）内网环境/外网环境：{network}")
    add_body(doc, f"（2）实施地点：{location}")
    add_body(doc, f"（3）专有云版本：{version}")

    add_body(doc, "（4）涉及的组件实例信息：")
    org = env.get("organization", "")
    resource_set = env.get("resource_set", "")
    add_body(doc, f"组织：{org}")
    add_body(doc, f"资源集：{resource_set}")

    # 实例信息（按操作类型分派：创建用表格，回收用文本行）
    table_spec = get_instance_table_spec(op_type)
    instances = data.get("instances", [])
    _build_instance_info(doc, op_type, instances, table_spec)

    # ================================================================
    # 4. 实施计划
    # ================================================================
    add_heading_1(doc, "4. 实施计划")

    plan = data.get("plan", {})

    add_heading_2(doc, "4.1 检修窗口")
    window = plan.get("window", {})
    if window:
        add_table(doc, ["年份", "开始时间", "结束时间"],
                  [[window.get("year", ""), window.get("start", ""), window.get("end", "")]])

    add_heading_2(doc, "4.2 实施人员")
    personnel = plan.get("personnel", {})
    if personnel:
        rows = [
            ["方案提供人", personnel.get("provider", "")],
            ["检修执行人", personnel.get("executor", "")],
            ["检修复核人", personnel.get("reviewer", "")],
            ["业务系统参与人", personnel.get("biz_participant", "不涉及")],
            ["安全责任人", personnel.get("security_officer", "")],
        ]
        add_table(doc, ["角色", "姓名"], rows)

    # ================================================================
    # 5. 风险评估
    # ================================================================
    add_heading_1(doc, "5. 风险评估")

    risk = data.get("risk_assessment", {})

    add_heading_2(doc, "5.1 影响范围")
    add_body(doc, risk.get("impact", f"{title}对业务无影响。"))

    add_heading_2(doc, "5.2 危险点分析")

    dangers = risk.get("dangers", [
        "（1）授权不当危险点：授权过大，导致操作影响预定方案以外的生产环境实例。",
        "（2）备份不当危险点：检修前实例未进行备份或备份无效，导致操作失败后无法完成应急回退。",
        "（3）验证不当危险点：操作对象以及服务是否存在单点隐患未核实清楚，导致操作后出现业务影响。",
        "（4）双人复核不当危险点：双人复核不仔细，导致操作错误执行而出现业务影响。",
    ])
    for d in dangers:
        add_body(doc, d)

    add_heading_2(doc, "5.3 安全措施")
    safety = risk.get("safety_measures", {})

    add_heading_3(doc, "5.3.1 授权")
    ascm_account = safety.get("ascm_account", "")
    bastion_accounts = safety.get("bastion_accounts", "")
    add_body(doc, f"ASCM：国网总部直属单位权限      授权账号：{ascm_account}")
    add_body(doc, f"堡垒机账号：{bastion_accounts}")

    add_heading_3(doc, "5.3.2 备份")
    backup_desc = safety.get("backup", get_default_backup(title, op_type))
    add_body(doc, f"（1）{backup_desc}")

    add_heading_3(doc, "5.3.3 验证")
    verification_desc = safety.get("verification", get_default_verification(title, op_type))
    add_body(doc, f"（1）{verification_desc}")

    add_heading_3(doc, "5.3.4 双人复核")
    add_body(doc, "（1）确认在正确的组织和资源集下做操作，检查实例操作对象是否正确；")
    add_body(doc, "（2）严格按照文档复核关键步骤及关键点。")

    # ================================================================
    # 6. 实施步骤
    # ================================================================
    add_heading_1(doc, "6. 实施步骤")

    steps = data.get("steps", {})

    add_heading_2(doc, "6.1 备份")
    backup_step = steps.get("backup", backup_desc)
    add_body(doc, f"1、{backup_step}")

    add_heading_2(doc, "6.2 检修前验证")
    pre_verify = steps.get("pre_verification", verification_desc)
    add_body(doc, f"1、{pre_verify}")

    add_heading_2(doc, "6.3 检修操作")
    op_steps = steps.get("operation", [])
    add_body(doc, _get_op_heading(op_type, title))

    # 重复实例信息（与第3节一致，按类型分派）
    _build_instance_info(doc, op_type, instances, table_spec)

    if op_steps:
        for i, step in enumerate(op_steps, 1):
            add_body(doc, f"{i}、{step}")
    else:
        _build_default_op_steps(doc, op_type, org, resource_set, ascm_account, instances)

    add_heading_2(doc, "6.4 检修后验证")
    post_verify = steps.get("post_verification",
                            f"验证{title}正常；联系项目组验证业务正常；")
    add_body(doc, f"1、{post_verify}")

    # ================================================================
    # 7. 回滚步骤
    # ================================================================
    add_heading_1(doc, "7. 回滚步骤")

    rollback = data.get("rollback", {})

    add_heading_2(doc, "7.1 回滚操作")
    rollback_op = rollback.get("operation", [])

    if rollback_op:
        add_body(doc, f"1、{title}回退")
        for i, step in enumerate(rollback_op, 1):
            add_body(doc, f"（{i}）{step}")
    else:
        _build_default_rollback(doc, op_type, title)

    add_heading_2(doc, "7.2 回滚后验证")
    rb_verify = rollback.get("verification",
                             f"验证{title}回退正常；联系项目组验证业务正常")
    add_body(doc, f"1、{rb_verify}")

    # 保存
    doc.save(output_path)
    print(f"检修方案已生成：{output_path}")
    return output_path


# ================================================================
# 入口
# ================================================================

def main():
    parser = argparse.ArgumentParser(description="生成Redis实例检修方案 .docx")
    parser.add_argument("--data", required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 文件路径")
    args = parser.parse_args()

    data_path = Path(args.data)
    if not data_path.exists():
        print(f"错误：数据文件不存在：{args.data}")
        sys.exit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    output_path = args.output
    if not output_path.endswith(".docx"):
        output_path += ".docx"

    generate_plan(data, output_path)


if __name__ == "__main__":
    main()
