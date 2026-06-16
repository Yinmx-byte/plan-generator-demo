"""Maintenance plan generation service."""

import json
import os
import re
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from fastapi import HTTPException

from agents.plan_agent import PlanAgentRuntime, run_plan_agent
from rag import get_knowledge_base
from runtime import (
    SKILLS_ROOT,
    build_system_prompt,
    get_agent_knowledge,
    get_formatter,
    get_model,
    get_skill_registry,
    get_skill_toolkit,
)
from scripts.generate_plan import build_document
from services.json_utils import extract_json, get_response_text

_generated_files: dict[str, Path] = {}


def compact_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").strip())


def derive_plan_title(state: dict[str, str], data: dict[str, Any]) -> str:
    """Keep the document title as a short plan name, not the raw demand."""
    model_title = compact_text(data.get("title") or data.get("document", {}).get("title", ""))
    demand_markers = ("需求", "需要", "项目组", "问题工单", "通过本次", "完成资源", "。", "，", ",")
    if model_title and len(model_title) <= 28 and not any(marker in model_title for marker in demand_markers):
        return model_title if model_title.endswith("检修方案") else f"{model_title}检修方案"

    network = compact_text(state.get("network"))
    source = compact_text(
        state.get("background")
        or state.get("instances")
        or state.get("maintenance_type")
        or model_title
        or "云平台"
    )
    subject = re.split(r"因|由于|为|需要|项目组|已提报|，|。|,", source, maxsplit=1)[0]
    subject = re.sub(r"(创建|新增|新建|申请|回收|释放|删除|调整|变更|重启|扩容|缩容|升配|降配).*$", "", subject)
    subject = subject or compact_text(state.get("instances") or state.get("maintenance_type") or "云平台")
    if network and network not in subject and network in {"内网", "外网", "内外网", "内、外网"}:
        subject = f"{network}{subject}"

    action = compact_text(state.get("maintenance_type") or "检修")
    title = f"{subject}{action}检修方案"
    title = re.sub(r"(检修方案)+$", "检修方案", title)
    return title[:40]


def get_plan_agent_runtime() -> PlanAgentRuntime:
    """Wire runtime dependencies for the maintenance plan composition agent."""
    return PlanAgentRuntime(
        build_system_prompt=build_system_prompt,
        get_model=get_model,
        get_formatter=get_formatter,
        get_toolkit=get_skill_toolkit,
        get_agent_knowledge=get_agent_knowledge,
        get_response_text=get_response_text,
    )


def get_generated_file(file_id: str | None) -> Optional[Path]:
    path = _generated_files.get(file_id or "")
    return path if path and path.exists() else None


def get_generated_path(session: dict[str, Any]) -> Optional[Path]:
    generated = session.get("generated") or {}
    return get_generated_file(generated.get("file_id"))


async def build_generation_orchestration_context(state: dict[str, str]) -> dict[str, Any]:
    rag_chunks: list[str] = []
    knowledge_base = get_knowledge_base(SKILLS_ROOT)
    if knowledge_base is not None:
        rag_query = f"""检修方案参考资料检索
检修类型：{state.get("maintenance_type", "")}
网络环境：{state.get("network", "")}
实施地点：{state.get("location", "")}
涉及实例：{state.get("instances", "")}
技术参数：{state.get("tech_params", "")}
补充要求：{state.get("ops_detail", "")}
请检索相似内部模板、阿里云通用检修方案、风险控制、前置检查、实施步骤、回退和验证要求。"""
        try:
            rag_chunks = await knowledge_base.retrieve(
                rag_query,
                top_k=int(os.getenv("PLAN_RAG_TOP_K", os.getenv("RAG_TOP_K", "5"))),
            )
        except Exception:
            rag_chunks = []

    if rag_chunks:
        blocks = []
        max_chars = int(os.getenv("PLAN_RAG_CONTEXT_MAX_CHARS", "6000"))
        used = 0
        for idx, chunk in enumerate(rag_chunks, start=1):
            clean = chunk.strip()
            if not clean:
                continue
            room = max_chars - used
            if room <= 0:
                break
            clean = clean[:room]
            used += len(clean)
            blocks.append(f"[RAG-{idx}]\n{clean}")
        rag_context = "\n\n".join(blocks) if blocks else "当前未检索到 RAG 参考资料。"
    else:
        rag_context = "当前未检索到 RAG 参考资料。请确认百炼知识库配置、远程文档解析状态、索引任务状态和检索阈值。"

    skill_prompt = get_skill_registry().get_agent_skill_prompt()
    return {
        "skill_selection_mode": "agentscope_react_auto",
        "rag_enabled": knowledge_base is not None,
        "rag_chunks_count": len(rag_chunks),
        "prompt_context": (
            "## 编排上下文：Skill 加载方式\n"
            "Skill 选择完全交给 AgentScope 注册信息和 ReActAgent 自主判断。所有 Skill 已通过 AgentScope Toolkit.register_agent_skill 注册，"
            "请根据 AgentScope 暴露的 Skill name/description 自主判断需要读取哪些 SKILL.md。\n\n"
            "## AgentScope Skill 摘要\n"
            f"{skill_prompt}\n\n"
            "## 编排上下文：RAG 参考资料\n"
            f"{rag_context}\n\n"
            "## 使用要求\n"
            "- Skill 是主规则来源：文档结构、必填章节、风险点、实施步骤和脚本模板优先遵循 ReActAgent 自主读取的 Skill。\n"
            "- RAG 是参考依据：用于补充内部模板措辞、历史方案经验、阿里云通用方案/API 约束，不得覆盖 Skill 的硬性规则。\n"
            "- 输出 JSON 中建议包含 evidence 字段，记录 agent_selected_skills 和 rag_chunks_count，便于后续审计。\n"
        ),
    }


async def repair_and_extract_json(text: str) -> dict:
    """Repair slightly malformed model JSON and parse it again."""
    try:
        return extract_json(text)
    except json.JSONDecodeError:
        repair_prompt = f"""下面是一段模型输出的检修方案 JSON，但它可能存在漏逗号、尾随逗号、代码块包裹或混入说明文字等格式问题。
请在不改变语义和字段内容的前提下，把它修复为严格合法 JSON。

要求：
- 只输出 JSON 对象
- 不要输出 markdown
- 不要解释
- 保留 document、sections、tables、steps 等原有结构

原始输出：
{text}
"""
        response = await get_model()(
            [
                {"role": "system", "content": "你是 JSON 修复器，只输出严格合法 JSON。"},
                {"role": "user", "content": repair_prompt},
            ],
        )
        try:
            return extract_json(get_response_text(response))
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=502,
                detail="模型返回的方案 JSON 无法解析，请重试或补充更明确的需求。",
            ) from exc


def write_model_output_debug(text: str, prefix: str = "plan_model_output") -> Path:
    output_dir = Path(tempfile.gettempdir()) / "plan-generator"
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    path.write_text(text or "", encoding="utf-8")
    return path


def build_fallback_plan_data(
    state: dict[str, str],
    orchestration: dict[str, Any],
    raw_text: str = "",
) -> dict[str, Any]:
    """Build a renderer-safe generic document when model JSON cannot be parsed."""
    maintenance_type = state.get("maintenance_type") or "检修"
    title = derive_plan_title(state, {"title": f"{maintenance_type}检修方案"})
    checked_type = maintenance_type.strip()
    type_names = ["配置变更", "组件升级", "组件扩缩容", "数据库变更", "日常维护（原硬件设备）", "其他"]
    checkbox_items = [
        {
            "label": name,
            "checked": name == checked_type or (name != "其他" and name in checked_type),
            "extra": checked_type if name == "其他" and checked_type not in type_names else "",
        }
        for name in type_names
    ]
    if not any(item["checked"] for item in checkbox_items):
        checkbox_items[-1]["checked"] = True
        checkbox_items[-1]["extra"] = checked_type or "待实施前确认"

    return {
        "title": title,
        "department": "云运营中心平台运维处",
        "date": datetime.now().strftime("%Y年%m月%d日"),
        "evidence": {
            "skill_selection_mode": orchestration.get("skill_selection_mode", "agentscope_react_auto"),
            "rag_enabled": orchestration.get("rag_enabled", False),
            "rag_chunks_count": orchestration.get("rag_chunks_count", 0),
            "fallback_used": True,
            "raw_output_saved": bool(raw_text),
        },
        "document": {
            "title": title,
            "cover": {"logo_width_cm": 3.1, "top_spacers": 7, "middle_spacers": 8},
            "header": [
                {"text": "云运营中心平台运维处", "font_size": 14, "align": "center"},
                {"text": datetime.now().strftime("%Y年%m月%d日"), "font_size": 12, "align": "center"},
            ],
            "sections": [
                {
                    "heading": "背景",
                    "level": 1,
                    "blocks": [
                        {"type": "paragraph", "text": state.get("background") or "根据业务运维需求，需要制定并执行本次检修方案。", "first_line_indent": 0.74},
                        {"type": "paragraph", "text": "该事项项目组提报问题工单，需检修进行处理，实现问题工单闭环。", "first_line_indent": 0.74},
                    ],
                },
                {"heading": "检修类型", "level": 1, "blocks": [{"type": "checkbox_group", "items": checkbox_items, "per_line": 3}]},
                {
                    "heading": "现场环境",
                    "level": 1,
                    "blocks": [
                        {"type": "paragraph", "text": f"（1）内网环境/外网环境：{state.get('network') or '待实施前确认'}"},
                        {"type": "paragraph", "text": f"（2）实施地点：{state.get('location') or '待实施前确认'}"},
                        {"type": "paragraph", "text": "（3）专有云版本：v3.16"},
                        {"type": "paragraph", "text": "（4）涉及的组件实例信息："},
                        {"type": "paragraph", "text": state.get("instances") or "待实施前确认", "first_line_indent": 0.74},
                    ],
                },
                {
                    "heading": "实施计划",
                    "level": 1,
                    "blocks": [
                        {"type": "heading", "text": "4.1 检修窗口", "level": 2},
                        {"type": "table", "columns": [{"key": "year", "label": "年份"}, {"key": "start_time", "label": "开始时间"}, {"key": "end_time", "label": "结束时间"}], "rows": [{"year": state.get("schedule_year", ""), "start_time": state.get("schedule_start", ""), "end_time": state.get("schedule_end", "")}]},
                        {"type": "heading", "text": "4.2 实施人员", "level": 2},
                        {"type": "table", "columns": [{"key": "provider", "label": "方案提供人"}, {"key": "executor", "label": "检修执行人"}, {"key": "reviewer", "label": "检修复核人"}, {"key": "business_participant", "label": "业务系统参与人"}, {"key": "security_officer", "label": "安全责任人"}], "rows": [{"provider": state.get("provider", ""), "executor": state.get("executor", ""), "reviewer": state.get("reviewer", ""), "business_participant": "待实施前确认", "security_officer": state.get("security_officer", "")}]},
                    ],
                },
                {
                    "heading": "风险评估",
                    "level": 1,
                    "blocks": [
                        {"type": "heading", "text": "5.1影响范围", "level": 2},
                        {"type": "paragraph", "text": "模型输出无法解析，本章节保留通用占位。请重新生成以读取产品 Skill 并补齐具体影响范围。", "first_line_indent": 0.74},
                        {"type": "heading", "text": "5.2危险点分析", "level": 2},
                        {"type": "paragraphs", "items": ["（1）授权不当危险点：授权过大，导致操作影响预定方案以外的生产环境实例。", "（2）备份不当危险点：检修前备份、快照、原配置或业务确认记录不完整，导致异常后无法准确恢复。", "（3）验证不当危险点：检修对象、业务访问状态、监控指标或项目组确认结果未核实清楚，导致操作后出现业务影响。", "（4）双人复核不当危险点：双人复核不仔细，导致操作错误执行而出现业务影响。"]},
                        {"type": "heading", "text": "5.3安全措施", "level": 2},
                        {"type": "heading", "text": "5.3.1授权", "level": 3},
                        {"type": "paragraphs", "items": [f"ASCM：国网总部直属单位权限     授权账号：{state.get('ascm_account') or '待实施前确认'}", f"堡垒机账号：{state.get('bastion_account') or '待实施前确认'}"]},
                        {"type": "heading", "text": "5.3.2备份", "level": 3},
                        {"type": "paragraph", "text": "按对应产品 Skill 和实际检修动作确认备份、快照、原配置导出或业务确认记录。", "first_line_indent": 0.74},
                        {"type": "heading", "text": "5.3.3验证", "level": 3},
                        {"type": "paragraph", "text": "按对应产品 Skill 验证检修前后资源状态、业务访问、监控告警和日志。", "first_line_indent": 0.74},
                        {"type": "heading", "text": "5.3.4 双人复核", "level": 3},
                        {"type": "paragraphs", "items": ["（1）确认在正确的组织和资源集下做操作，检查操作对象是否正确；", "（2）严格按照文档复核关键步骤及关键点。"]},
                    ],
                },
                {
                    "heading": "实施步骤",
                    "level": 1,
                    "blocks": [
                        {"type": "heading", "text": "6.1备份", "level": 2},
                        {"type": "numbered_list", "items": ["按产品 Skill 要求确认备份、快照、配置导出或业务确认记录。"]},
                        {"type": "heading", "text": "6.2 检修前验证", "level": 2},
                        {"type": "numbered_list", "items": ["确认检修对象、窗口期、授权账号、审批工单、资源状态和监控均满足实施条件。"]},
                        {"type": "heading", "text": "6.3 检修操作", "level": 2},
                        {"type": "heading", "text": f"6.3.1 {state.get('instances') or maintenance_type}", "level": 3},
                        {"type": "numbered_list", "items": ["模型输出无法解析，未能可靠生成产品级操作步骤。请重新生成或补充更明确需求。"]},
                        {"type": "heading", "text": "6.4 检修后验证", "level": 2},
                        {"type": "numbered_list", "items": ["验证资源状态、业务访问、监控告警和日志，确认无异常后关闭检修。"]},
                    ],
                },
                {
                    "heading": "回滚步骤",
                    "level": 1,
                    "blocks": [
                        {"type": "heading", "text": "7.1 回滚操作", "level": 2},
                        {"type": "numbered_list", "items": ["触发回滚条件时立即停止后续操作，并按产品 Skill、备份或原配置执行恢复。"]},
                        {"type": "heading", "text": "7.2 回滚后验证", "level": 2},
                        {"type": "numbered_list", "items": ["回滚后重新验证业务访问、资源状态、监控告警和日志，并保留记录。"]},
                    ],
                },
            ],
        },
    }

def build_user_prompt(
    background: str, maintenance_type: str, network: str, location: str,
    instances: str, schedule_year: str, schedule_start: str, schedule_end: str,
    provider: str, executor: str, reviewer: str, security_officer: str,
    ascm_account: str, bastion_account: str, ops_detail: str, tech_params: str,
    orchestration_context: str = "",
    edit_instruction: str = "",
    previous_document_text: str = "",
) -> str:
    """Convert form fields into the user task for the plan-writing ReActAgent."""
    edit_context = ""
    if edit_instruction:
        previous_text = previous_document_text[:8000] if previous_document_text else "No previous document text was available."
        edit_context = f"""
## Document revision task
This is a revision of the previously generated maintenance plan, not a brand-new plan.
Keep the existing structure and unchanged content as much as possible. Only update the
personnel, risks, steps, rollback content, scripts, or other parts explicitly requested.

## Revision request
{edit_instruction}

## Previous document text for reference
{previous_text}
"""
    return f"""请根据以下检修需求生成标准化检修方案 JSON。

后端已经通过 AgentScope 注册全部可用 Skill，并提供 RAG 检索上下文。你必须根据 Skill 摘要自主判断需要哪些 Skill，并使用 read_file 工具读取/核对相关 SKILL.md；若涉及多个检修类型，应组合多个 Skill。

输出结构硬性要求：
- 必须读取并遵守 `maintenance-plan-composer` Skill 中的“Word 渲染 JSON 格式契约”。
- 必须读取并遵守具体产品检修 Skill；产品 Skill 决定风险项、操作步骤、回滚步骤和参数表。
- 生成最终答案前，调用 `build_maintenance_document` 工具对完整 JSON 做一次渲染检查；工具通过后仍然只输出完整 JSON。
- 顶层必须包含 document 对象；每个章节必须包含 blocks 数组。
- 表格、标题、正文、复选框、编号步骤必须按该 Skill 规定的 block 类型输出。
- 不要只输出 heading/title/content；只有标题没有 blocks 的章节会被判定为无效格式。
- 检修方案必须至少包含实施计划表或参数表，不允许全部正文都是普通段落。
- 后端不会再为 ECS、SLB、RDS 等产品补写专用步骤；如果你没有按产品 Skill 生成可执行步骤，最终文档会缺失关键内容。

{orchestration_context}
{edit_context}

## 背景与检修事项
{background}

## 检修类型
{maintenance_type}

## 现场环境
内/外网环境：{network}
实施地点：{location}
涉及的组件实例描述：
{instances}

## 检修窗口
{schedule_year} {schedule_start} 至 {schedule_end}

## 人员信息
方案提供人：{provider}
检修执行人：{executor}
检修复核人：{reviewer}
安全责任人：{security_officer}

## 授权账号
ASCM账号：{ascm_account}
堡垒机账号：{bastion_account}

## 检修操作补充说明（可选）
{ops_detail}

注意：即使补充说明为空，也必须根据已激活 Skill 自动生成完整的“六、实施步骤”，不要要求用户手工提供详细操作步骤。

## 技术参数
{tech_params}
"""


def count_document_body_blocks(data: dict[str, Any]) -> int:
    spec = data.get("document")
    if not isinstance(spec, dict):
        return 1

    def count_blocks(section: Any) -> int:
        if isinstance(section, str):
            return 1 if section.strip() else 0
        if not isinstance(section, dict):
            return 0
        total = 0
        blocks = section.get("blocks")
        if isinstance(blocks, list):
            for block in blocks:
                if isinstance(block, str) and block.strip():
                    total += 1
                elif isinstance(block, dict):
                    if block.get("text") or block.get("content") or block.get("description"):
                        total += 1
                    if block.get("items"):
                        total += len(block.get("items") or [])
                    if block.get("rows"):
                        total += len(block.get("rows") or [])
        for key in ("content", "body", "text", "description"):
            value = section.get(key)
            if isinstance(value, str) and value.strip():
                total += 1
            elif isinstance(value, list):
                total += len([item for item in value if str(item).strip()])
        for key in ("items", "points", "steps"):
            value = section.get(key)
            if isinstance(value, list):
                total += len([item for item in value if str(item).strip()])
        for child in section.get("children") or section.get("subsections") or []:
            total += count_blocks(child)
        return total

    sections = spec.get("sections") or spec.get("chapters") or []
    return sum(count_blocks(section) for section in sections)


def document_format_stats(data: dict[str, Any]) -> dict[str, int]:
    """Count canonical render blocks before accepting a model document spec."""
    spec = data.get("document")
    stats = {"typed_blocks": 0, "tables": 0, "sections": 0}
    if not isinstance(spec, dict):
        return stats

    def visit_section(section: Any) -> None:
        if not isinstance(section, dict):
            return
        stats["sections"] += 1
        for block in section.get("blocks") or []:
            if not isinstance(block, dict) or "type" not in block:
                continue
            stats["typed_blocks"] += 1
            if block.get("type") == "table":
                stats["tables"] += 1
        for child in section.get("children") or section.get("subsections") or []:
            visit_section(child)

    for section in spec.get("sections") or spec.get("chapters") or []:
        visit_section(section)
    return stats


REQUIRED_SECTION_NAMES = ["背景", "检修类型", "现场环境", "实施计划", "风险评估", "实施步骤", "回滚步骤"]


def normalize_heading(value: Any) -> str:
    text = re.sub(r"\s+", "", str(value or ""))
    text = re.sub(r"^[一二三四五六七八九十]+、", "", text)
    return text.lower()


def default_section(name: str, state: dict[str, str]) -> dict[str, Any]:
    if name == "背景":
        return {
            "heading": name,
            "level": 1,
            "blocks": [
                {"type": "paragraph", "text": state.get("background") or "待根据需求和对应 Skill 补充。", "first_line_indent": 0.74}
            ],
        }
    if name == "检修类型":
        maintenance_type = state.get("maintenance_type") or "其他"
        items = [
            {"label": item, "checked": item in maintenance_type, "extra": ""}
            for item in ["配置变更", "组件升级", "组件扩缩容", "数据库变更", "日常维护（原硬件设备）", "其他"]
        ]
        if not any(item["checked"] for item in items):
            items[-1]["checked"] = True
            items[-1]["extra"] = maintenance_type
        return {"heading": name, "level": 1, "blocks": [{"type": "checkbox_group", "items": items, "per_line": 3}]}
    if name == "现场环境":
        return {
            "heading": name,
            "level": 1,
            "blocks": [
                {"type": "paragraph", "text": f"（1）内网环境/外网环境：{state.get('network') or '待实施前确认'}"},
                {"type": "paragraph", "text": f"（2）实施地点：{state.get('location') or '待实施前确认'}"},
                {"type": "paragraph", "text": "（3）专有云版本：v3.16"},
                {"type": "paragraph", "text": "（4）涉及的组件实例信息："},
                {"type": "paragraph", "text": state.get("instances") or "待实施前确认", "first_line_indent": 0.74},
            ],
        }
    if name == "实施计划":
        return {
            "heading": name,
            "level": 1,
            "blocks": [
                {"type": "heading", "text": "4.1 检修窗口", "level": 2},
                {"type": "table", "columns": [{"key": "year", "label": "年份"}, {"key": "start_time", "label": "开始时间"}, {"key": "end_time", "label": "结束时间"}], "rows": [{"year": state.get("schedule_year", ""), "start_time": state.get("schedule_start", ""), "end_time": state.get("schedule_end", "")}]},
                {"type": "heading", "text": "4.2 实施人员", "level": 2},
                {"type": "table", "columns": [{"key": "provider", "label": "方案提供人"}, {"key": "executor", "label": "检修执行人"}, {"key": "reviewer", "label": "检修复核人"}, {"key": "business_participant", "label": "业务系统参与人"}, {"key": "security_officer", "label": "安全责任人"}], "rows": [{"provider": state.get("provider", ""), "executor": state.get("executor", ""), "reviewer": state.get("reviewer", ""), "business_participant": "待实施前确认", "security_officer": state.get("security_officer", "")}]},
            ],
        }
    return {
        "heading": name,
        "level": 1,
        "blocks": [
            {
                "type": "paragraph",
                "text": "模型未输出本章节内容，请重新生成或补充更明确需求。",
                "first_line_indent": 0.74,
            }
        ],
    }


def ensure_renderer_ready_document(
    data: dict[str, Any],
    state: dict[str, str],
    orchestration: dict[str, Any],
) -> dict[str, Any]:
    """Apply generic renderer guards without encoding product-specific logic."""
    spec = data.setdefault("document", {})
    if not isinstance(spec, dict):
        data["document"] = spec = {}

    sections = spec.get("sections")
    if not isinstance(sections, list):
        sections = []

    normalized_sections: list[dict[str, Any]] = []
    existing_by_name: dict[str, dict[str, Any]] = {}
    for section in sections:
        if not isinstance(section, dict):
            continue
        heading = str(section.get("heading") or section.get("title") or "").strip()
        if not heading:
            continue
        section["heading"] = heading
        section.setdefault("level", 1)
        blocks = section.get("blocks")
        section["blocks"] = blocks if isinstance(blocks, list) and blocks else []
        normalized_sections.append(section)
        existing_by_name.setdefault(normalize_heading(heading), section)

    for required_name in REQUIRED_SECTION_NAMES:
        if normalize_heading(required_name) not in existing_by_name:
            normalized_sections.append(default_section(required_name, state))

    spec["sections"] = normalized_sections
    data.setdefault("evidence", {})
    if isinstance(data["evidence"], dict):
        data["evidence"]["renderer_schema_guard"] = True
        data["evidence"].setdefault(
            "skill_selection_mode",
            orchestration.get("skill_selection_mode", "agentscope_react_auto"),
        )
    return data

def docx_to_markdown(raw: bytes) -> str:
    from docx import Document

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(raw)
        tmp_path = Path(tmp.name)
    try:
        document = Document(str(tmp_path))
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
                lines.append("")
        for table_idx, table in enumerate(document.tables, start=1):
            lines.append(f"表格 {table_idx}")
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if rows:
                width = max(len(row) for row in rows)
                normalized = [row + [""] * (width - len(row)) for row in rows]
                lines.append("| " + " | ".join(normalized[0]) + " |")
                lines.append("| " + " | ".join(["---"] * width) + " |")
                for row in normalized[1:]:
                    lines.append("| " + " | ".join(row) + " |")
                lines.append("")
        return "\n".join(lines).strip() + "\n"
    finally:
        tmp_path.unlink(missing_ok=True)


async def generate_docx_from_state(
    state: dict[str, str],
    orchestration: Optional[dict[str, Any]] = None,
    trace_callback=None,
    edit_instruction: str = "",
    previous_document_text: str = "",
) -> tuple[str, Path, str]:
    if orchestration is None:
        orchestration = await build_generation_orchestration_context(state)
    data = await compose_plan_json(
        state,
        orchestration,
        trace_callback=trace_callback,
        edit_instruction=edit_instruction,
        previous_document_text=previous_document_text,
    )
    data = validate_plan_json(data, state, orchestration)
    return render_docx(data)


async def compose_plan_json(
    state: dict[str, str],
    orchestration: dict[str, Any],
    trace_callback=None,
    edit_instruction: str = "",
    previous_document_text: str = "",
) -> dict[str, Any]:
    """Compose structured maintenance-plan JSON with the Plan ReActAgent.

    The ReActAgent remains the document-writing specialist because it can load
    AgentScope Skills and call the render-check tool. This service boundary
    keeps the rest of the backend coupled to a stable JSON-composition API
    instead of to the agent implementation.
    """
    user_prompt = build_user_prompt(
        **state,
        orchestration_context=orchestration["prompt_context"],
        edit_instruction=edit_instruction,
        previous_document_text=previous_document_text,
    )
    response = await run_plan_agent(user_prompt, get_plan_agent_runtime(), trace_callback=trace_callback)
    text = get_response_text(response)
    try:
        data = await repair_and_extract_json(text)
    except HTTPException:
        write_model_output_debug(text)
        data = build_fallback_plan_data(state, orchestration, text)
    return data


def collect_plan_validation_issues(data: dict[str, Any]) -> list[str]:
    """Collect structural issues that matter for an actionable plan."""
    issues: list[str] = []
    spec = data.get("document")
    if not isinstance(spec, dict):
        return ["missing_document_object"]
    sections = spec.get("sections") or spec.get("chapters") or []
    if not isinstance(sections, list) or not sections:
        return ["missing_document_sections"]

    normalized_headings = [
        normalize_heading(section.get("heading") or section.get("title"))
        for section in sections
        if isinstance(section, dict)
    ]
    for required_name in REQUIRED_SECTION_NAMES:
        required = normalize_heading(required_name)
        if not any(required in heading or heading in required for heading in normalized_headings):
            issues.append(f"missing_section:{required_name}")

    def section_block_count(keyword: str) -> int:
        count = 0
        for section in sections:
            if not isinstance(section, dict):
                continue
            heading = normalize_heading(section.get("heading") or section.get("title"))
            if keyword not in heading:
                continue
            count += len(section.get("blocks") or [])
            for block in section.get("blocks") or []:
                if isinstance(block, dict):
                    count += len(block.get("items") or [])
                    count += len(block.get("rows") or [])
        return count

    if section_block_count("风险") < 4:
        issues.append("risk_section_too_thin")
    if section_block_count("实施步骤") < 4:
        issues.append("implementation_steps_too_thin")
    return issues


def validate_plan_json(
    data: dict[str, Any],
    state: dict[str, str],
    orchestration: dict[str, Any],
) -> dict[str, Any]:
    """Validate and normalize structured plan JSON before DOCX rendering."""
    validation_issues = collect_plan_validation_issues(data)
    format_stats = document_format_stats(data)
    fallback_reason = ""
    if count_document_body_blocks(data) < 3:
        fallback_reason = "model_document_had_too_few_body_blocks"
    elif isinstance(data.get("document"), dict) and format_stats["typed_blocks"] < 8:
        fallback_reason = "model_document_was_not_canonical_template_format"

    if fallback_reason:
        debug_path = write_model_output_debug(
            json.dumps(data, ensure_ascii=False, indent=2),
            prefix="invalid_plan_model_output",
        )
        data = build_fallback_plan_data(state, orchestration, json.dumps(data, ensure_ascii=False))
        data.setdefault("evidence", {})
        if isinstance(data["evidence"], dict):
            data["evidence"]["fallback_reason"] = fallback_reason
            data["evidence"]["debug_path"] = str(debug_path)

    data.setdefault("department", "云运营中心平台运维处")
    data.setdefault("date", datetime.now().strftime("%Y年%m月%d日"))
    data.setdefault("evidence", {})
    if isinstance(data["evidence"], dict):
        data["evidence"].setdefault("rag_enabled", orchestration["rag_enabled"])
        data["evidence"].setdefault(
            "rag_chunks_count",
            orchestration["rag_chunks_count"],
        )
        data["evidence"].setdefault(
            "skill_selection_mode",
            orchestration.get("skill_selection_mode", "agentscope_react_auto"),
        )
        data["evidence"]["validation_issues"] = validation_issues

    data = ensure_renderer_ready_document(data, state, orchestration)
    normalized_title = derive_plan_title(state, data)
    data["title"] = normalized_title
    document_spec = data.setdefault("document", {})
    if isinstance(document_spec, dict):
        document_spec["title"] = normalized_title
    return data


def render_docx(data: dict[str, Any]) -> tuple[str, Path, str]:
    """Render validated plan JSON into a Word document and register download."""
    doc = build_document(data)
    output_dir = Path(tempfile.gettempdir()) / "plan-generator"
    output_dir.mkdir(parents=True, exist_ok=True)
    file_id = uuid.uuid4().hex
    output_path = output_dir / f"检修方案_{file_id[:8]}.docx"
    doc.save(str(output_path))
    _generated_files[file_id] = output_path

    files = sorted(
        output_dir.glob("检修方案_*.docx"),
        key=lambda f: f.stat().st_mtime,
        reverse=True,
    )
    for old_file in files[10:]:
        try:
            old_file.unlink()
        except OSError:
            pass

    filename = data.get("title") or data.get("document", {}).get("title", "检修方案")
    return file_id, output_path, filename + ".docx"




